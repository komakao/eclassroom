from django.shortcuts import render, redirect
from teacher.models import *
from student.models import *
from student.lesson import *
from account.avatar import *
from django.views import generic
from django.contrib.auth.models import User, Group
from django.views.generic import CreateView, UpdateView, DeleteView, ListView, RedirectView, TemplateView
from teacher.forms import *
from account.forms import PasswordForm, NicknameForm
import django_excel as excel
import xlsxwriter
from django.http import HttpResponse, HttpResponseRedirect
from django.contrib.auth.decorators import login_required, user_passes_test
from django.db.models import Q
from django.http import JsonResponse
from django.core.exceptions import ObjectDoesNotExist, MultipleObjectsReturned
from collections import OrderedDict
from django.core.files.storage import FileSystemStorage
from uuid import uuid4
from docx import *
from docx.shared import Inches
from docx.shared import RGBColor
from io import BytesIO
import io
from django.utils.timezone import localtime
from django.utils.decorators import method_decorator
import json
from django.utils.http import urlquote
from django.conf import settings
from wsgiref.util import FileWrapper
from io import BytesIO

def filename_browser(request, filename):
	browser = request.META['HTTP_USER_AGENT'].lower()
	if 'edge' in browser:
		response['Content-Disposition'] = 'attachment; filename='+urlquote(filename)+'; filename*=UTF-8\'\'' + urlquote(filename)
		return response			
	elif 'webkit' in browser:
		# Safari 3.0 and Chrome 2.0 accepts UTF-8 encoded string directly.
		filename_header = 'filename=%s' % filename.encode('utf-8').decode('ISO-8859-1')
	elif 'trident' in browser or 'msie' in browser:
		# IE does not support internationalized filename at all.
		# It can only recognize internationalized URL, so we do the trick via routing rules.
		filename_header = 'filename='+filename.encode("BIG5").decode("ISO-8859-1")					
	else:
		# For others like Firefox, we follow RFC2231 (encoding extension in HTTP headers).
		filename_header = 'filename*="utf8\'\'' + str(filename.encode('utf-8').decode('ISO-8859-1')) + '"'
	return filename_header		

# 判斷是否為同班同學
def is_classmate(user, classroom_id):
    enroll_pool = [enroll for enroll in Enroll.objects.filter(classroom_id=classroom_id).order_by('seat')]
    student_ids = map(lambda a: a.student_id, enroll_pool)
    if user.id in student_ids:
        return True
    else:
        return False	

# 判斷是否為授課教師
def is_teacher(user, classroom_id):
    if Classroom.objects.filter(teacher_id=user.id, id=classroom_id).exists():
        return True
    elif Assistant.objects.filter(user_id=user.id, classroom_id=classroom_id).exists():
        return True
    return False

def in_teacher_group(user):
    if not user.groups.filter(name='teacher').exists():
        if not Assistant.objects.filter(user_id=user.id).exists():
            return False
    return True
	
	
class ClassroomTeacherRequiredMixin(object):	
    def dispatch(self, request, *args, **kwargs):
        if 'classroom_id' in kwargs:
            classroom_id = self.kwargs['classroom_id']
        else:
            classroom_id = self.kwargs['pk']
        user = self.request.user
        jsonDec = json.decoder.JSONDecoder()	
        classroom_list = []
        profile = Profile.objects.get(user=user)
        if len(profile.classroom) > 0 :		
            classroom_list = jsonDec.decode(profile.classroom)
        if str(classroom_id) in classroom_list:
            if not user.groups.filter(name='teacher').exists():
                if not Assistant.objects.filter(user_id=user.id).exists():		
                    return redirect("/")
            return super(ClassroomTeacherRequiredMixin, self).dispatch(request,*args, **kwargs)
        else :
            return redirect('/')
			
class TeacherRequiredMixin(object):
    def dispatch(self, request, *args, **kwargs):
        user = self.request.user
        if not user.groups.filter(name='teacher').exists():
            if not Assistant.objects.filter(user_id=user.id).exists():		
                return redirect("/")
        return super(TeacherRequiredMixin, self).dispatch(request,
            *args, **kwargs)	

class ClassroomList(TeacherRequiredMixin, generic.ListView):
    model = Classroom
    ordering = ['-id']
    paginate_by = 30
	
    def get_queryset(self):
        classrooms = Classroom.objects.filter(teacher_id=self.request.user.id).order_by("-id")
        return classrooms
    
class ClassroomCreate(TeacherRequiredMixin, CreateView):
    model =Classroom
    fields = ["lesson", "name", "password", "progress", "online"]
    success_url = "/teacher/classroom"   
    template_name = 'form.html'
      
    def form_valid(self, form):
        valid = super(ClassroomCreate, self).form_valid(form)
        classroom = form.save(commit=False)
        classroom.teacher_id = self.request.user.id
        classroom.save() 
        enroll = Enroll(classroom_id=classroom.id, student_id=classroom.teacher_id, seat=0)
        enroll.save()
        try :
            group = Group.objects.get(name="class"+str(classroom.id))	
        except ObjectDoesNotExist :
            group = Group(name="class"+str(classroom.id))
            group.save()     
        group.user_set.add(self.request.user)	
        jsonDec = json.decoder.JSONDecoder()	
        classroom_list = []
        user = self.request.user
        profile = Profile.objects.get(user=user)
        if len(profile.classroom) > 0 :		
            classroom_list = jsonDec.decode(profile.classroom)
        classroom_list.append(str(classroom.id))
        profile.classroom = json.dumps(classroom_list)
        profile.save()		
        # 指定作業分組
        lesson = classroom.lesson
        queryset = []		
        works = Work.objects.filter(user_id=self.request.user.id, lesson=lesson).order_by("id")		
        if lesson == 1 :
            for assignment in lesson_list1:
                queryset.append(assignment)
            for assignment in queryset:
                workgroup = WorkGroup(classroom_id=classroom.id, index=assignment[3])
                workgroup.save()
        else:
            for unit in lesson_list[int(lesson)-2][1]:
                for assignment in unit[1]:
                    queryset.append(assignment)     
            for assignment in queryset:
                workgroup = WorkGroup(classroom_id=classroom.id, index=assignment[2])
                workgroup.save()
        return valid
    
class ClassroomUpdate(ClassroomTeacherRequiredMixin, UpdateView):
    model = Classroom
    fields = ["lesson", "name", "password", "progress", "online"]
    success_url = "/teacher/classroom"   
    template_name = 'form.html'
	
    def dispatch(self, request, *args, **kwargs):
        user = self.request.user
        jsonDec = json.decoder.JSONDecoder()	
        classroom_list = []
        profile = Profile.objects.get(user=user)
        if len(profile.classroom) > 0 :		
            classroom_list = jsonDec.decode(profile.classroom)
        if str(self.kwargs['pk']) in classroom_list:
            return super(ClassroomUpdate, self).dispatch(request,*args, **kwargs)
        else :
            return redirect('/account/login/0')
	
# 教師可以查看所有帳號
class StudentListView(TeacherRequiredMixin, ListView):
    context_object_name = 'users'
    paginate_by = 50
    template_name = 'teacher/student_list.html'

    def get_queryset(self):
        username = username__icontains=self.request.user.username+"_"
        if self.request.GET.get('account') != None:
            keyword = self.request.GET.get('account')
            queryset = User.objects.filter(Q(username__icontains=username+keyword) | (Q(first_name__icontains=keyword) & Q(username__icontains=username))).order_by('-id')
        else :
            queryset = User.objects.filter(username__icontains=username).order_by('-id')
        return queryset

    def get_context_data(self, **kwargs):
        context = super(StudentListView, self).get_context_data(**kwargs)
        account = self.request.GET.get('account')
        context.update({'account': account})
        return context
	
# 教師可以查看所有帳號
class StudentJoinView(TeacherRequiredMixin, ListView):
    context_object_name = 'users'
    paginate_by = 60
    template_name = 'teacher/student_join.html'

    def get_queryset(self):
        enrolls = Enroll.objects.filter(classroom_id=self.kwargs['classroom_id'])	
        user_ids = [enroll.student_id for enroll in enrolls]
        username = username__icontains=self.request.user.username+"_"
        if self.request.GET.get('account') != None:
            keyword = self.request.GET.get('account')
            queryset = User.objects.filter((~Q(id__in=user_ids)) & (Q(username__icontains=username+keyword) | (Q(first_name__icontains=keyword)) & Q(username__icontains=username))).order_by('-id')
        else :
            queryset = User.objects.filter(~Q(id__in=user_ids) & Q(username__icontains=username)).order_by('-id')
        return queryset

    def get_context_data(self, **kwargs):
        context = super(StudentJoinView, self).get_context_data(**kwargs)
        account = self.request.GET.get('account')
        context.update({'account': account})
        context['classroom_id']=self.kwargs['classroom_id']
        return context


#加選學生
class StudentEnrollView(TeacherRequiredMixin, RedirectView):

    def get(self, request, *args, **kwargs):
        classroom_id = self.kwargs['classroom_id'] 
        classroom = Classroom.objects.get(id=classroom_id)
        students = self.request.POST.getlist('student')
        jsonDec = json.decoder.JSONDecoder()	
        classroom_list = []		
        for student in students:
            student_id = student.split(":")[0]
            student_seat = student.split(":")[1] 
            student_computer = student.split(":")[2]                         
            enroll = Enroll(student_id=student_id, classroom_id=classroom_id, seat=student_seat, computer=student_computer)
            enroll.save()
            user = User.objects.get(id=student_id)
            profile = Profile.objects.get(user=user)
            if len(profile.classroom) > 0 :		
                classroom_list = jsonDec.decode(profile.classroom)
            classroom_list.append(str(self.kwargs['classroom_id']))
            profile.classroom = json.dumps(classroom_list)
            profile.save()			
            messages = Message.objects.filter(author_id=classroom.teacher_id, classroom_id=classroom_id, type=1)
            for message in messages:
                try:
                    messagepoll = MessagePoll.objects.get(message_type=1, message_id=message.id, reader_id=student_id, classroom_id=classroom_id)
                except ObjectDoesNotExist:
                    messagepoll = MessagePoll(message_type=1, message_id=message.id, reader_id=student_id, classroom_id=classroom_id)
                    messagepoll.save()	
                except MultipleObjectsReturned:
                    pass				
        return super(StudentEnrollView, self).get(self, request, *args, **kwargs)        
        
    def get_redirect_url(self, *args, **kwargs):
        #TaxRate.objects.get(id=int(kwargs['pk'])).delete()   
        return '/student/classroom/'+ str(self.kwargs['classroom_id']) + '/classmate'
		
# Create your views here.
def import_sheet(request):
    if not in_teacher_group(request.user):
        return redirect("/")
    if request.method == "POST":
        form = UploadFileForm(request.POST,
                              request.FILES)
        if form.is_valid():
            ImportUser.objects.all().delete()
            request.FILES['file'].save_to_database(
                name_columns_by_row=0,
                model=ImportUser,
                mapdict=['username', 'password'])
            users = ImportUser.objects.all()
            return render(request, 'teacher/import_student.html',{'users':users})
        else:
            return HttpResponseBadRequest()
    else:
        form = UploadFileForm()
    return render(
        request,
        'teacher/import_form.html',
        {
            'form': form,
            'title': 'Excel file upload and download example',
            'header': ('Please choose any excel file ' +
                       'from your cloned repository:')
        })

# Create your views here.
def import_student(request):
    if not in_teacher_group(request.user):
        return redirect("/")

    users = ImportUser.objects.all()
    username_list = [request.user.username+"_"+user.username for user in users]
    exist_users = [user.username for user in User.objects.filter(username__in=username_list)]
    create_list = []
    for user in users:
        username = request.user.username+"_"+user.username
        if username in exist_users:
            continue
        new_user = User(username=username, first_name=user.username, password=user.password, email=username+"@edu.tw")
        new_user.set_password(user.password)
        create_list.append(new_user)

    User.objects.bulk_create(create_list)
    new_users = User.objects.filter(username__in=[user.username for user in create_list])

    profile_list = []
    message_list = []
    poll_list = []
    title = "請洽詢任課教師課程名稱及選課密碼"
    url = "/student/classroom/join"
    message = Message(title=title, url=url, time=timezone.now())
    message.save()
    for user in new_users:
        profile = Profile(user=user)
        profile_list.append(profile)
        poll = MessagePoll(message_id=message.id, reader_id=user.id)
        poll_list.append(poll)

    Profile.objects.bulk_create(profile_list)
    MessagePoll.objects.bulk_create(poll_list)

    return redirect('/teacher/student/list')

# 修改密碼
def password(request, user_id):
    if not in_teacher_group(request.user):
        return redirect("/")
    user = User.objects.get(id=user_id)
    if not user.username.startswith(request.user.username):
        return redirect("/")
    if request.method == 'POST':
        form = PasswordForm(request.POST)
        if form.is_valid():
            user = User.objects.get(id=user_id)
            user.set_password(request.POST['password'])
            user.save()
            return redirect('/teacher/student/list/')
    else:
        form = PasswordForm()
        user = User.objects.get(id=user_id)

    return render(request, 'form.html',{'form': form, 'user':user})

# 修改暱稱
def nickname(request, user_id):
    if not in_teacher_group(request.user):
        return redirect("/")
    user = User.objects.get(id=user_id)
    if not user.username.startswith(request.user.username):
        return redirect("/")
    if request.method == 'POST':
        form = NicknameForm(request.POST)
        if form.is_valid():
            user = User.objects.get(id=user_id)
            user.first_name =form.cleaned_data['first_name']
            user.save()
            return redirect('/teacher/student/list/')
    else:
        user = User.objects.get(id=user_id)
        form = NicknameForm(instance=user)
    return render(request, 'form.html',{'form': form})
			
# 設定班級助教
@login_required
def classroom_assistant(request, classroom_id):
    if not is_teacher(request.user, classroom_id):
        return redirect("/")
		
    assistants = Assistant.objects.filter(classroom_id=classroom_id).order_by("-id")
    classroom = Classroom.objects.get(id=classroom_id)

    return render(request, 'teacher/assistant.html',{'assistants': assistants, 'classroom':classroom})

# 教師可以查看所有帳號
class AssistantListView(ClassroomTeacherRequiredMixin, ListView):
    context_object_name = 'users'
    paginate_by = 20
    template_name = 'teacher/assistant_user.html'

    def get_queryset(self):
        if self.request.GET.get('account') != None:
            keyword = self.request.GET.get('account')
            if self.kwargs['group'] == 1:
                queryset = User.objects.filter(Q(groups__name='apply') & (Q(username__icontains=keyword) | Q(first_name__icontains=keyword))).order_by("-id")
            else :
                queryset = User.objects.filter(Q(username__icontains=keyword) | Q(first_name__icontains=keyword)).order_by('-id')		
        else :
            if self.kwargs['group'] == 1:
                queryset = User.objects.filter(groups__name='apply').order_by("-id")
            else :
                queryset = User.objects.all().order_by('-id')
        return queryset

    def get_context_data(self, **kwargs):
        context = super(AssistantListView, self).get_context_data(**kwargs)
        context['classroom'] = Classroom.objects.get(id=self.kwargs['classroom_id'])
        context['group'] = self.kwargs['group']
        assistant_list = []
        assistants = Assistant.objects.filter(classroom_id=self.kwargs['classroom_id'])
        for assistant in assistants:
            assistant_list.append(assistant.user_id)
        context['assistants'] = assistant_list
        return context

# 列出所有助教課程
class AssistantClassroomListView(TeacherRequiredMixin, ListView):
    model = Classroom
    context_object_name = 'classrooms'
    template_name = 'teacher/assistant_list.html'
    paginate_by = 20

    def get_queryset(self):
        assistants = Assistant.objects.filter(user_id=self.request.user.id)
        classroom_list = []
        for assistant in assistants:
            classroom_list.append(assistant.classroom_id)
        queryset = Classroom.objects.filter(id__in=classroom_list).order_by("-id")
        return queryset

# Ajax 設為助教、取消助教
def assistant_make(request):
    if not in_teacher_group(request.user):
        return JsonResponse({'status':'fail'}, safe=False)
		
    classroom_id = request.POST.get('classroomid')
    user_id = request.POST.get('userid')
    action = request.POST.get('action')
    user = User.objects.get(id=user_id)
    if user_id and action :
        if action == 'set':
            try :
                assistant = Assistant.objects.get(classroom_id=classroom_id, user_id=user_id)
            except ObjectDoesNotExist :
                assistant = Assistant(classroom_id=classroom_id, user_id=user_id)
                assistant.save()
            # 將助教設為0號學生
            enrolls = Enroll.objects.filter(classroom_id=classroom_id, student_id=user_id)
            if len(enrolls) == 0:
                enroll = Enroll(classroom_id=classroom_id, student_id=user_id, seat=0)
                enroll.save()
            try :
                group = Group.objects.get(name="class"+classroom_id)	
            except ObjectDoesNotExist :
                group = Group(name="class"+classroom_id)
                group.save()     
            group.user_set.add(request.user)	
            jsonDec = json.decoder.JSONDecoder()	
            classroom_list = []
            profile = Profile.objects.get(user=user)
            if len(profile.classroom) > 0 :		
                classroom_list = jsonDec.decode(profile.classroom)
            classroom_list.append(str(classroom_id))
            profile.classroom = json.dumps(classroom_list)
            profile.save()	
        else :
            try :
                assistant = Assistant.objects.get(classroom_id=classroom_id, user_id=user_id)
                assistant.delete()
                enroll = Enroll.objects.filter(classroom_id=classroom_id, student_id=user_id)
                enroll.delete()
            except ObjectDoesNotExist :
                pass
            try :
                group = Group.objects.get(name="class"+classroom_id)	
            except ObjectDoesNotExist :
                group = Group(name="class"+classroom_id)
                group.save()     
            group.user_set.remove(request.user)
            jsonDec = json.decoder.JSONDecoder()	
            classroom_list = []
            profile = Profile.objects.get(user=user)
            if len(profile.classroom) > 0 :		
                classroom_list = jsonDec.decode(profile.classroom)
                classroom_list.remove(str(classroom_id))
            profile.classroom = json.dumps(classroom_list)
            profile.save()				
        return JsonResponse({'status':'ok'}, safe=False)
    else:
        return JsonResponse({'status':'fail'}, safe=False)
		
	
# 列出所有討論主題
class ForumListView(ClassroomTeacherRequiredMixin, ListView):
    model = FWork
    context_object_name = 'forums'
    template_name = "teacher/forum_list.html"		
    paginate_by = 20
			
    def get_queryset(self):        
        fclasses = FClass.objects.filter(classroom_id=self.kwargs['classroom_id']).order_by("-publication_date", "-forum_id")
        forums = []
        for fclass in fclasses:
            forum = FWork.objects.get(id=fclass.forum_id)
            forums.append([forum, fclass])
        return forums
			
    def get_context_data(self, **kwargs):
        context = super(ForumListView, self).get_context_data(**kwargs)
        classroom = Classroom.objects.get(id=self.kwargs['classroom_id'])
        context['classroom'] = classroom
        return context	
        
#新增一個討論主題
class ForumCreateView(ClassroomTeacherRequiredMixin, CreateView):
    model = FWork
    form_class = ForumForm
    template_name = "teacher/forum_form.html"
    def form_valid(self, form):
        self.object = form.save(commit=False)
        self.object.teacher_id = self.request.user.id
        self.object.classroom_id = self.kwargs['classroom_id']     
        self.object.save()  
        classrooms = self.request.POST.getlist('classrooms')
        for classroom in classrooms:
          forum_class = FClass(forum_id=self.object.id, classroom_id=classroom)
          forum_class.save()
        
        return redirect("/teacher/forum/"+str(self.kwargs['classroom_id']))           
        
    def get_context_data(self, **kwargs):
        context = super(ForumCreateView, self).get_context_data(**kwargs)
        classroom_list = []
        classrooms = Classroom.objects.filter(teacher_id=self.request.user.id)
        for classroom in classrooms:
            classroom_list.append(classroom.id)
        assistants = Assistant.objects.filter(user_id=self.request.user.id)
        for assistant in assistants:
            if not assistant.classroom_id in classroom_list:
                classroom_list.append(assistant.classroom_id)
        classrooms = Classroom.objects.filter(id__in=classroom_list).order_by("-id")
        context['classrooms'] = classrooms
        context['classroom_id'] = int(self.kwargs['classroom_id'])
        context['classroom'] = Classroom.objects.get(id=self.kwargs['classroom_id'])
        return context	
  
        return redirect("/teacher/forum/"+self.kwargs['classroom_id'])        
	
# 列出所有討論主題
class ForumAllListView(ClassroomTeacherRequiredMixin, ListView):
    model = FWork
    context_object_name = 'forums'
    template_name = "teacher/forum_all.html"		
    paginate_by = 20

    def get_queryset(self):
      queryset = FWork.objects.all().order_by("-id")
      if self.request.GET.get('account') != None:
        keyword = self.request.GET.get('account')
        users = User.objects.filter(Q(username__icontains=keyword) | Q(first_name__icontains=keyword)).order_by('-id')
        user_list = []
        for user in users:
            user_list.append(user.id)
        forums = queryset.filter(teacher_id__in=user_list)
        return forums
      else:				
        return queryset
			
    def get_context_data(self, **kwargs):
        context = super(ForumAllListView, self).get_context_data(**kwargs)
        context['categroy'] = self.kwargs['categroy']							
        context['categroy_id'] = self.kwargs['categroy_id']							
        return context	

# 展示討論素材
def forum_show(request, clssroom_id, forum_id):
    if not is_teacher(request.user, classroom_id):
        return redirect("/")
    forum = FWork.objects.get(id=forum_id)
    contents = FContent.objects.filter(forum_id=forum_id)
    return render_to_response('teacher/forum_show.html',{'contents':contents, 'forum':forum}, context_instance=RequestContext(request))

		
# 列出某討論主題的班級
class ForumClassList(ClassroomTeacherRequiredMixin, ListView):
    model = FWork
    context_object_name = 'classrooms'
    template_name = "teacher/forum_class.html"		
    paginate_by = 20
	
    def get_queryset(self):        		
        fclass_dict = dict(((fclass.classroom_id, fclass) for fclass in FClass.objects.filter(forum_id=self.kwargs['forum_id'])))		
        classroom_list = []
        classroom_ids = []
        classrooms = Classroom.objects.filter(teacher_id=self.request.user.id).order_by("-id")
        for classroom in classrooms:
            if classroom.id in fclass_dict:
                classroom_list.append([classroom, True, fclass_dict[classroom.id].deadline, fclass_dict[classroom.id].deadline_date])
            else :
                classroom_list.append([classroom, False, False, timezone.now()])
            classroom_ids.append(classroom.id)
        assistants = Assistant.objects.filter(user_id=self.request.user.id)
        for assistant in assistants:
            classroom = Classroom.objects.get(id=assistant.classroom_id)
            if not classroom.id in classroom_ids:
                if classroom.id in fclass_dict:
                    classroom_list.append([classroom, True, fclass_dict[classroom.id].deadline, fclass_dict[classroom.id].deadline_date])
                else :
                    classroom_list.append([classroom, False, False, timezone.now()])
        return classroom_list
			
    def get_context_data(self, **kwargs):
        context = super(ForumClassList, self).get_context_data(**kwargs)				
        fwork = FWork.objects.get(id=self.kwargs['forum_id'])
        context['fwork'] = fwork
        context['forum_id'] = self.kwargs['forum_id']
        return context	
	
# Ajax 開放班取、關閉班級
def forum_switch(request):
    if not in_teacher_group(request.user):
        return JsonResponse({'status':status}, safe=False)      
    forum_id = request.POST.get('forumid')
    classroom_id = request.POST.get('classroomid')		
    status = request.POST.get('status')
    try:
        fwork = FClass.objects.get(forum_id=forum_id, classroom_id=classroom_id)
        if status == 'false' :
    	      fwork.delete()
    except ObjectDoesNotExist:
        if status == 'true':
            fwork = FClass(forum_id=forum_id, classroom_id=classroom_id)
            fwork.save()
    return JsonResponse({'status':status}, safe=False)        
	
# 列出所有討論主題素材
class ForumContentList(ClassroomTeacherRequiredMixin, ListView):
    model = FContent
    context_object_name = 'contents'
    template_name = "teacher/forum_content.html"		
    def get_queryset(self):
        queryset = FContent.objects.filter(forum_id=self.kwargs['forum_id']).order_by("-id")
        return queryset
			
    def get_context_data(self, **kwargs):
        context = super(ForumContentList, self).get_context_data(**kwargs)
        fwork = FWork.objects.get(id=self.kwargs['forum_id'])
        fclasses = FClass.objects.filter(forum_id=self.kwargs['forum_id'])				
        context['fwork']= fwork
        context['forum_id'] = self.kwargs['forum_id']
        context['fclasses'] = fclasses
        context['classroom_id'] = self.kwargs['classroom_id']		
        return context	
			
#新增一個課程
class ForumContentCreate(ClassroomTeacherRequiredMixin, CreateView):
    model = FContent
    form_class = ForumContentForm
    template_name = "teacher/forum_content_form.html"
    def form_valid(self, form):
        self.object = form.save(commit=False)
        work = FContent(forum_id=self.object.forum_id)
        if self.object.types == 1:
            work.types = 1
            work.title = self.object.title
            work.link = self.object.link
        if self.object.types  == 2:
            work.types = 2					
            work.youtube = self.object.youtube
        if self.object.types  == 3:
            work.types = 3
            myfile = self.request.FILES['content_file']
            fs = FileSystemStorage()
            filename = uuid4().hex
            work.title = myfile.name
            work.filename = str(self.request.user.id)+"/"+filename
            fs.save("static/upload/"+str(self.request.user.id)+"/"+filename, myfile)
        if self.object.types  == 4:
            work.types = 4
        work.memo = self.object.memo
        work.save()         
  
        return redirect("/teacher/forum/content/"+str(self.kwargs['classroom_id'])+"/"+str(self.kwargs['forum_id']))

    def get_context_data(self, **kwargs):
        context = super(ForumContentCreate, self).get_context_data(**kwargs)
        context['forum'] = FWork.objects.get(id=self.kwargs['forum_id'])
        return context

def forum_delete(request, classroom_id, forum_id, content_id):
    if not is_teacher(request.user, classroom_id):
        return redirect("/")
    instance = FContent.objects.get(id=content_id)
    instance.delete()

    return redirect("/teacher/forum/content/"+str(classroom_id)+str(forum_id))  
	
def forum_edit(request, classroom_id, forum_id, content_id):
    if not is_teacher(request.user, classroom_id):
        return redirect("/")
    try:
        instance = FContent.objects.get(id=content_id)
    except:
        pass
    if request.method == 'POST':
            content_id = request.POST.get("id", "")
            try:
                content = FContent.objects.get(id=content_id)
            except ObjectDoesNotExist:
	              content = FContent(forum_id= request.POST.get("forum_id", ""), types=form.cleaned_data['types'])
            if content.types == 1:
                content.title = request.POST.get("title", "")
                content.link = request.POST.get("link", "")
            elif content.types == 2:
                content.youtube = request.POST.get("youtube", "")
            elif content.types == 3:
                myfile =  request.FILES.get("content_file", "")
                fs = FileSystemStorage()
                filename = uuid4().hex
                content.title = myfile.name
                content.filename = str(request.user.id)+"/"+filename
                fs.save("static/upload/"+str(request.user.id)+"/"+filename, myfile)
            content.memo = request.POST.get("memo", "")
            content.save()
            return redirect('/teacher/forum/content/'+str(classroom_id)+str(forum_id))   
    return render(request, 'teacher/forum_edit.html',{'content': instance, 'forum_id':forum_id, 'content_id':content_id})		
	
def forum_download(request, classroom_id, content_id):
    if not is_teacher(request.user, classroom_id):
        return redirect("/")
    content = FContent.objects.get(id=content_id)
    filename = content.title
    download =  settings.BASE_DIR + "/static/upload/" + content.filename
    wrapper = FileWrapper(open( download, "rb" ))
    response = HttpResponse(wrapper, content_type = 'application/force-download')
    #response = HttpResponse(content_type='application/force-download')
    filename_header = filename_browser(request, filename)
    response['Content-Disposition'] = 'attachment; ' + filename_header	
    # It's usually a good idea to set the 'Content-Length' header too.
    # You can also set any other required headers: Cache-Control, etc.
    return response
    #return render_to_response('student/download.html', {'download':download})
		
class ForumEditUpdate(ClassroomTeacherRequiredMixin, UpdateView):
    model = FWork
    fields = ["title"]        
    template_name = 'teacher/forum_form.html'
    
    def get_success_url(self):
        succ_url =  '/teacher/forum/'+str(self.kwargs['classroom_id'])
        return succ_url       

    def get_context_data(self):
        context = super(ForumEditUpdate, self).get_context_data()        
        context['classroom'] = Classroom.objects.get(id=self.kwargs['classroom_id'])        
        context['classroom_id'] =self.kwargs['classroom_id']
        context['classrooms'] = Classroom.objects.filter(teacher_id=self.request.user.id).order_by("-id")
        return context        
      

def forum_export(request, classroom_id, forum_id):
	if not is_teacher(request.user, classroom_id):
		return redirect("/")
	classroom = Classroom.objects.get(id=classroom_id)
	try:
		fwork = FWork.objects.get(id=forum_id)
		enrolls = Enroll.objects.filter(classroom_id=classroom_id)
		datas = []
		contents = FContent.objects.filter(forum_id=forum_id).order_by("-id")
		fwork = FWork.objects.get(id=forum_id)
		works_pool = SFWork.objects.filter(index=forum_id).order_by("-id")
		reply_pool = SFReply.objects.filter(index=forum_id).order_by("-id")	
		file_pool = SFContent.objects.filter(index=forum_id, visible=True).order_by("-id")	
		for enroll in enrolls:
			works = list(filter(lambda w: w.student_id==enroll.student_id, works_pool))
			if len(works)>0:
				replys = list(filter(lambda w: w.work_id==works[0].id, reply_pool))
			else:
				replys = []
			files = list(filter(lambda w: w.student_id==enroll.student_id, file_pool))
			if enroll.seat > 0:
				datas.append([enroll, works, replys, files])
		def getKey(custom):
			return -custom[0].seat
		datas = sorted(datas, key=getKey, reverse=True)	
		#word
		document = Document()
		docx_title=u"討論區-" + classroom.name + "-"+ str(timezone.localtime(timezone.now()).date())+".docx"
		document.add_paragraph(request.user.first_name + u'的討論區作業')
		document.add_paragraph(u'主題：'+fwork.title)		
		document.add_paragraph(u"班級：" + classroom.name)		
		
		for enroll, works, replys, files in datas:
			user = User.objects.get(id=enroll.student_id)
			run = document.add_paragraph().add_run(str(enroll.seat)+")"+user.first_name)
			font = run.font
			font.color.rgb = RGBColor(0xFA, 0x24, 0x00)
			if len(works)>0:
				#p = document.add_paragraph(str(works[0].publication_date)[:19]+'\n'+works[0].memo)
				p = document.add_paragraph(str(localtime(works[0].publication_date))[:19]+'\n')
				# 將 memo 以時間標記為切割點，切分為一堆 tokens
				tokens = re.split('(\[m_\d+#\d+:\d+:\d+\])', works[0].memo)
				# 依續比對 token 格式
				for token in tokens:
					m = re.match('\[m_(\d+)#(\d+):(\d+):(\d+)\]', token)
					if m: # 若為時間標記，則插入連結
						vid = filter(lambda material: material.id == int(m.group(1)), contents)[0]
						add_hyperlink(document, p, vid.youtube+"&t="+m.group(2)+"h"+m.group(3)+"m"+m.group(4)+"s", "["+m.group(2)+":"+m.group(3)+":"+m.group(4)+"]")
					else: # 以一般文字插入
						p.add_run(token)
			if len(replys)>0:
				for reply in replys:
					user = User.objects.get(id=reply.user_id)
					run = document.add_paragraph().add_run(user.first_name+u'>'+str(localtime(reply.publication_date))[:19]+u'>留言:\n'+reply.memo)
					font = run.font
					font.color.rgb = RGBColor(0x42, 0x24, 0xE9)		
			if len(files)>0:
				for file in files:
					if file.visible:
						if file.title[-3:].upper() == "PNG" or file.title[-3:].upper() == "JPG":
							filename = 'static/upload/'+file.filename
							if os.path.exists(filename):
								copyfile(filename, 'static/upload/file.png')					
								document.add_picture('static/upload/file.png',width=Inches(6.0))
						else:
							p = document.add_paragraph()
							full_url = request.build_absolute_uri()
							index = full_url.find("/",9)
							url = full_url[:index] + "/student/forum/download/" + str(file.id) 
							add_hyperlink(document, p, url, file.title)
		# Prepare document for download        
		f = io.BytesIO()
		document.save(f)
		length = f.tell()
		f.seek(0)
		response = HttpResponse(
			f.getvalue(),
			content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
		)

		filename_header = filename_browser(request, docx_title)
		response['Content-Disposition'] = 'attachment; ' + filename_header
		response['Content-Length'] = length
		return response

	except ObjectDoesNotExist:
		pass
	return True		


def add_hyperlink(document, paragraph, url, name):
    """
    Add a hyperlink to a paragraph.
    :param document: The Document being edited.
    :param paragraph: The Paragraph the hyperlink is being added to.
    :param url: The url to be added to the link.
    :param name: The text for the link to be displayed in the paragraph
    :return: None
    """

    part = document.part
    rId = part.relate_to(url, RT.HYPERLINK, is_external=True)

    init_hyper = OxmlElement('w:hyperlink')
    init_hyper.set(qn('r:id'), rId, )
    init_hyper.set(qn('w:history'), '1')

    new_run = OxmlElement('w:r')

    rPr = OxmlElement('w:rPr')

    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')

    rPr.append(rStyle)
    new_run.append(rPr)
    new_run.text = name
    init_hyper.append(new_run)

    r = paragraph.add_run()
    r._r.append(init_hyper)
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return None

def forum_grade(request, classroom_id, action):
	if not is_teacher(request.user, classroom_id):
		return redirect("/")
	classroom = Classroom.objects.get(id=classroom_id)
	forum_ids = []
	forums = []
	fclasses = FClass.objects.filter(classroom_id=classroom_id).order_by("publication_date", "forum_id")
	for fclass in fclasses:
		forum_ids.append(fclass.forum_id)
		forum = FWork.objects.get(id=fclass.forum_id)
		forums.append(forum.title)
	enrolls = Enroll.objects.filter(classroom_id=classroom_id).order_by("seat")
	datas = {}
	for enroll in enrolls:
			sfworks = SFWork.objects.filter(index__in=forum_ids, student_id=enroll.student_id).order_by("id")
			if len(sfworks) > 0:
				for fclass in fclasses:
						works = list(filter(lambda w: w.index==fclass.forum_id, sfworks))
						if enroll.student_id in datas:
							if len(works) > 0 :
								datas[enroll.student_id].append(works[0])
							else :
								datas[enroll.student_id].append(SFWork())
						else:
							if len(works) > 0:
								datas[enroll.student_id] = [works[0]]
							else :
								datas[enroll.student_id] = [SFWork()]
			else :
				datas[enroll.student_id] = [SFWork()]
	results = []
	for enroll in enrolls:
		student_name = User.objects.get(id=enroll.student_id).first_name
		results.append([enroll, student_name, datas[enroll.student_id]])
	
	#下載Excel
	if action == 1:
		classroom = Classroom.objects.get(id=classroom_id)       
		output = BytesIO()
		workbook = xlsxwriter.Workbook(output)    
		worksheet = workbook.add_worksheet(classroom.name)
		date_format = workbook.add_format({'num_format': 'yy/mm/dd'})
		
		row = 1
		worksheet.write(row, 1, u'座號')
		worksheet.write(row, 2, u'姓名')
		index = 3
		for forum in forums:
			worksheet.write(row, index, forum)
			index += 1
		
		row += 1
		index = 3
		for fclass in fclasses:
			worksheet.write(row, index, datetime.strptime(str(fclass.publication_date)[:19],'%Y-%m-%d %H:%M:%S'), date_format)
			index += 1			

		for enroll, student_name, works in results:
			row += 1
			worksheet.write(row, 1, enroll.seat)
			worksheet.write(row, 2, student_name)
			index = 3
			for work in works:
				if work.id:
					worksheet.write(row, index, work.score)
				else:
					worksheet.write(row, index, '')
				index +=1 

		workbook.close()
		# xlsx_data contains the Excel file
		response = HttpResponse(content_type='application/vnd.ms-excel')
		filename = u'討論區-' + classroom.name + "-" + str(localtime(timezone.now()).date()) + '.xlsx'
		response['Content-Disposition'] = filename_browser(request, filename)
		xlsx_data = output.getvalue()
		response.write(xlsx_data)
		return response
	else :
		return render(request,'teacher/forum_grade.html',{'results':results, 'forums':forums, 'classroom_id':classroom_id, 'fclasses':fclasses})

class ForumDeadlineUpdate(ClassroomTeacherRequiredMixin, UpdateView):
    model = FWork
    fields = ['title']
    template_name = "teacher/forum_deadline_form.html"		   
	
    def form_valid(self, form):
        if is_teacher(self.request.user, self.kwargs['classroom_id']) or is_assistant(self.request.user, self.kwargs['classroom_id']):
            #fclass = .objects.get(id=self.kwargs['pk'])
            reduce = group.numbers - form.cleaned_data['numbers']
            if reduce > 0:
                for i in range(reduce):
                    StudentGroup.objects.filter(group_id=self.kwargs['pk'], group=group.numbers-i).delete()
            form.save()
        return HttpResponseRedirect(self.get_success_url())   
      
    def get_context_data(self):
        context = super(ForumDeadlineUpdate, self).get_context_data()
        classroom_id = self.kwargs['classroom_id']
        forum_id = self.kwargs['pk']
        context['classroom_id'] = classroom_id
        context['fclass'] = FClass.objects.get(classroom_id=classroom_id, forum_id=forum_id)
        context['fclasses'] = FClass.objects.filter(forum_id=forum_id)
        return context    
		
class ForumPublishReject(ClassroomTeacherRequiredMixin, RedirectView):
    def get_redirect_url(self, *args, **kwargs):
      index = self.kwargs['index']
      classroom_id = self.kwargs['classroom_id']
      user_id = self.kwargs['user_id']
      try:
          fwork = FWork.objects.get(id=index)
          works = SFWork.objects.filter(index=index, student_id=user_id).order_by("-id")
          work = works[0]
          work.publish = False
          work.save()
          update_avatar(user_id, 4, -2)
          # History
          history = PointHistory(user_id=user_id, kind=4, message=u'-2分--退回討論區作業<'+fwork.title+'>', url='/student/forum/memo/'+str(classroom_id)+'/'+str(index)+'/0')
          history.save()								
      except ObjectDoesNotExist:
            pass
      return "/student/forum/memo/"+str(classroom_id)+"/"+str(index)+"/0"
 
		
# Ajax 設定期限、取消期限
def forum_deadline_set(request):
    if not in_teacher_group(request.user):
        return JsonResponse({'status':status}, safe=False) 
    forum_id = request.POST.get('forumid')
    classroom_id = request.POST.get('classroomid')		
    status = request.POST.get('status')
    try:
        fclass = FClass.objects.get(forum_id=forum_id, classroom_id=classroom_id)
    except ObjectDoesNotExist:
        fclass = Fclass(forum_id=forum_id, classroom_id=classroom_id)
    if status == 'True':
        fclass.deadline = True
    else :
        fclass.deadline = False
    fclass.save()
    return JsonResponse({'status':status}, safe=False)        

# Ajax 設定期限日期
def forum_deadline_date(request):
    if not in_teacher_group(request.user):
        return JsonResponse({'status':deadline_date}, safe=False)  		
    forum_id = request.POST.get('forumid')
    classroom_id = request.POST.get('classroomid')		
    deadline_date = request.POST.get('deadlinedate')
    try:
        fclass = FClass.objects.get(forum_id=forum_id, classroom_id=classroom_id)
    except ObjectDoesNotExist:
        fclass = FClass(forum_id=forum_id, classroom_id=classroom_id)
    #fclass.deadline_date = deadline_date.strftime('%d/%m/%Y')
    fclass.deadline_date = datetime.strptime(deadline_date, '%Y/%m/%d %H:%M')
    fclass.save()
    return JsonResponse({'status':deadline_date}, safe=False)                  
					
		
# 列出所有課程
class WorkListView2(ClassroomTeacherRequiredMixin, ListView):
    model = TWork
    context_object_name = 'works'
    template_name = 'teacher/twork_list.html'
    paginate_by = 20

    def get_queryset(self):
        queryset = TWork.objects.filter(classroom_id=self.kwargs['classroom_id']).order_by("-id")
        return queryset

    def get_context_data(self, **kwargs):
        context = super(WorkListView2, self).get_context_data(**kwargs)
        context['classroom'] = Classroom.objects.get(id=self.kwargs['classroom_id'])
        return context

#新增一個課程
class WorkCreateView2(ClassroomTeacherRequiredMixin, CreateView):
    model = TWork
    form_class = WorkForm
    template_name = 'form.html'
    def form_valid(self, form):
        self.object = form.save(commit=False)
        self.object.teacher_id = self.request.user.id
        self.object.classroom_id = self.kwargs['classroom_id']
        self.object.save()
        workgroup = WorkGroup(typing=1, classroom_id=self.kwargs['classroom_id'], index=self.object.id)
        workgroup.save()
        return redirect("/teacher/work2/"+str(self.kwargs['classroom_id']))

# 修改選課密碼
def work_edit(request, classroom_id):
    if not is_teacher(request.user, classroom_id):
        return redirect("/")
    classroom = Classroom.objects.get(id=classroom_id)
    if request.method == 'POST':
        form = ClassroomForm(request.POST)
        if form.is_valid():
            classroom.name =form.cleaned_data['name']
            classroom.password = form.cleaned_data['password']
            classroom.save()
            return redirect('/teacher/classroom')
    else:
        form = ClassroomForm(instance=classroom)

    return render(request, 'form.html',{'form': form})

# 列出某作業所有同學名單
def work_class2(request, classroom_id, work_id):
    if not is_teacher(request.user, classroom_id):
        return redirect("/")
    enrolls = Enroll.objects.filter(classroom_id=classroom_id)
    classroom = Classroom.objects.get(id=classroom_id)
    lesson = classroom.lesson
    classmate_work = []
    scorer_name = ""
    for enroll in enrolls:
        try:
            work = Work.objects.get(typing=1, user_id=enroll.student_id, index=work_id, lesson=lesson)
            if work.scorer > 0 :
                scorer = User.objects.get(id=work.scorer)
                scorer_name = scorer.first_name
            else :
                scorer_name = "1"
        except ObjectDoesNotExist:
            work = Work(typing=1, index=work_id, user_id=0, lesson=lesson)
        except MultipleObjectsReturned:
            work = Work.objects.filter(typing=1, user_id=enroll.student_id, index=work_id, lesson=lesson).last()
        try:
            group_id = WorkGroup.objects.get(typing=1, classroom_id=classroom_id, index=work_id).group_id
        except ObjectDoesNotExist:
            group_id = 0
        if group_id > 0:
            try: 
                group = StudentGroup.objects.get(enroll_id=enroll.id, group_id=group_id).group
            except ObjectDoesNotExist:
                group = -1
        else :
            group = -1
        assistant = WorkAssistant.objects.filter(typing=1, group_id=group_id, student_id=enroll.student_id, lesson=lesson, index=work_id)
        if assistant.exists():
            classmate_work.append([enroll,work,1, scorer_name, group])
        else :
            classmate_work.append([enroll,work,0, scorer_name, group])
    def getKey(custom):
        return custom[0].seat

    classmate_work = sorted(classmate_work, key=getKey)
    try:
        group_id = WorkGroup.objects.get(typing=1, classroom_id=classroom_id, index=work_id).group_id
    except ObjectDoesNotExist:
        group_id = 0               

    return render(request, 'teacher/work_class.html',{'group_id': group_id, 'typing':1, 'classmate_work': classmate_work, 'classroom':classroom, 'index': work_id})

def work_list(request, classroom_id):
    if not is_teacher(request.user, classroom_id):
        return redirect("/")
    classroom = Classroom.objects.get(id=classroom_id)        
    group_id = classroom.group
    GroupModelFormset = modelformset_factory(WorkGroup, fields=['id', 'classroom_id','index', 'group_id', 'typing'],extra=0)	
    if request.method == 'POST':
        formset = GroupModelFormset(request.POST)
        if formset.is_valid():
            for form in formset :
                workgroup = WorkGroup.objects.get(classroom_id=classroom_id, index=form.cleaned_data['index'], typing=0)
                workgroup.group_id = form.cleaned_data['group_id']
                workgroup.save()
                form.save()
            return redirect('/teacher/work/'+str(classroom_id))
        else:
            return redirect("/")
    else:
        workgroups = WorkGroup.objects.filter(classroom_id=classroom_id, typing=0).order_by("index")
        groups = ClassroomGroup.objects.filter(classroom_id=classroom_id)
        formset = GroupModelFormset(queryset=workgroups)		
    return render(request, 'teacher/work_list.html', {'formset': formset, 'classroom':classroom, 'groups': groups})		
	
def work_groupset(request, typing, classroom_id):	
    if request.method == 'POST':
        WorkGroup.objects.filter(classroom_id=classroom_id, typing=typing).update(group_id=request.POST['group'])
        if typing == 0:
            return redirect("/teacher/work/"+str(classroom_id))
        else :
            return redirect("/teacher/work2/"+str(classroom_id))
    return redirect("/")		
	

def work_list2(request, classroom_id):
    if not is_teacher(request.user, classroom_id):
        return redirect("/")
    classroom = Classroom.objects.get(id=classroom_id)        
    group_id = classroom.group
    GroupModelFormset = modelformset_factory(WorkGroup, fields=['id', 'classroom_id','index', 'group_id', 'typing'],extra=0)	
    if request.method == 'POST':
        formset = GroupModelFormset(request.POST)
        if formset.is_valid():
            for form in formset :
                workgroup = WorkGroup.objects.get(typing=1, classroom_id=classroom_id, index=form.cleaned_data['index'])
                workgroup.group_id = form.cleaned_data['group_id']
                workgroup.save()
                form.save()
            return redirect('/teacher/work2/'+str(classroom_id))
        else:
            return redirect("/1")
    else:
        workgroups = WorkGroup.objects.filter(classroom_id=classroom_id, typing=1).order_by("index")
        groups = ClassroomGroup.objects.filter(classroom_id=classroom_id)
        formset = GroupModelFormset(queryset=workgroups)		
    return render(request, 'teacher/twork_list.html', {'formset': formset, 'classroom':classroom, 'groups': groups})		
	

# 列出某作業所有同學名單
def work_class(request, typing, classroom_id, index):
    if not is_teacher(request.user, classroom_id):
        return redirect("/")
		
    enrolls = Enroll.objects.filter(classroom_id=classroom_id)
    classroom = Classroom.objects.get(id=classroom_id)
    classmate_work = []
    scorer_name = ""
    try:
        group_id = WorkGroup.objects.get(typing=typing, classroom_id=classroom_id, index=index).group_id
    except ObjectDoesNotExist:
        group_id = 0

    for enroll in enrolls:
        try:
            work = Work.objects.get(typing=typing, user_id=enroll.student_id, index=index, lesson=classroom.lesson)
            if work.scorer > 0 :
                scorer = User.objects.get(id=work.scorer)
                scorer_name = scorer.first_name
            else :
                scorer_name = "1"
        except ObjectDoesNotExist:
            work = Work(typing=0, index=index, user_id=0, lesson=classroom.lesson)
        except MultipleObjectsReturned:
            work = Work.objects.filter(typing=0, user_id=enroll.student_id, index=index, lesson=classroom.lesson).last()
        if group_id > 0:
            try: 
                group = StudentGroup.objects.get(enroll_id=enroll.id, group_id=group_id).group
            except ObjectDoesNotExist:
                group = -1
        else :
            group = -1            
        assistant = WorkAssistant.objects.filter(typing=0, group=group, student_id=enroll.student_id, lesson=classroom.lesson, index=index)
        if assistant.exists():
            classmate_work.append([enroll,work,1, scorer_name, group])
        else :
            classmate_work.append([enroll,work,0, scorer_name, group])        

    def getKey(custom):
        return custom[0].seat

    classmate_work = sorted(classmate_work, key=getKey)
    return render(request, 'teacher/work_class.html',{'group_id':group_id, 'typing':0, 'classmate_work': classmate_work, 'classroom':classroom, 'index': index})

# 教師評分
@login_required
# 列出某作業所有同學名單
def work_group(request, typing, classroom_id, index):
    if not is_teacher(request.user, classroom_id):
        return redirect("/")
    enrolls = Enroll.objects.filter(classroom_id=classroom_id)
    classroom = Classroom.objects.get(id=classroom_id)
    lesson = classroom.lesson
    classmate_work = []
    scorer_name = ""
    for enroll in enrolls:
        try:
            work = Work.objects.get(typing=typing, user_id=enroll.student_id, index=index, lesson=classroom.lesson)
            if work.scorer > 0 :
                scorer = User.objects.get(id=work.scorer)
                scorer_name = scorer.first_name
            else :
                scorer_name = "1"
        except ObjectDoesNotExist:
            work = Work(typing=typing, index=index, user_id=0, lesson=classroom.lesson)
        except MultipleObjectsReturned:
            work = Work.objects.filter(typing=typing, user_id=enroll.student_id, index=index, lesson=classroom.lesson).last()
        try:
            group_id = WorkGroup.objects.get(typing=typing, classroom_id=classroom_id, index=index).group_id
        except ObjectDoesNotExist:
            group_id = 0
        if group_id > 0:
            try: 
                group = StudentGroup.objects.get(enroll_id=enroll.id, group_id=group_id).group
            except ObjectDoesNotExist:
                group = -1
        else :
            group = -1
        assistant = WorkAssistant.objects.filter(typing=typing, group=group, student_id=enroll.student_id, lesson=classroom.lesson, index=index)
        if assistant.exists():
            classmate_work.append([enroll,work,1, scorer_name, group, group_id])
        else :
            classmate_work.append([enroll,work,0, scorer_name, group, group_id])
    def getKey(custom):
        return custom[4], custom[1].publication_date
      
    classmate_work = sorted(classmate_work, key=getKey)    
    return render(request, 'teacher/work_group.html',{'test': group_id, 'typing':typing, 'classmate_work': classmate_work, 'classroom':classroom, 'index': index, 'lesson':lesson})
  
  
# (小)教師評分
def scoring(request, classroom_id, user_id, index, typing):
    user = User.objects.get(id=request.user.id)
    classroom = Classroom.objects.get(id=classroom_id)
    lesson = classroom.lesson
    teacher = is_teacher(user, classroom_id)	
    lesson_name = ""
    try:
        group_id = WorkGroup.objects.get(typing=typing, classroom_id=classroom_id, index=index).group_id
    except ObjectDoesNotExist:
        group_id = 0
    enroll_id = Enroll.objects.get(classroom_id=classroom_id, student_id=user_id).id
    try :
        studentgroup = StudentGroup.objects.get(enroll_id=enroll_id, group_id=group_id)
        group = studentgroup.group                                           
    except ObjectDoesNotExist:
        group = -1  

    if typing == 0:
        lesson_dict = OrderedDict()
        for unit1 in lesson_list[int(lesson)-1][1]:
            for assignment in unit1[1]:
                lesson_dict[assignment[2]] = assignment[0]
        lesson_name = lesson_dict[int(index)]
    elif typing == 1:
        lesson_name = TWork.objects.get(id=index).title
    user = User.objects.get(id=user_id)
    enroll = Enroll.objects.get(classroom_id=classroom_id, student_id=user_id)
    workfiles = []
    try:
        assistant = WorkAssistant.objects.filter(typing=typing, group=group, lesson=lesson, index=index,student_id=request.user.id)
    except ObjectDoesNotExist:
        if not is_teacher(request.user, classroom_id) or not is_assistant(request.user, classroom_id):
            return render(request, 'message.html', {'message':"您沒有權限"})
    try:
        work3 = Work.objects.get(typing=typing, user_id=user_id, index=index, lesson=lesson)
        pic = work3.id
    except ObjectDoesNotExist:
        work3 = Work(typing=typing, index=index, user_id=user_id, lesson=lesson)
        pic = 0
    except MultipleObjectsReturned:
        works = Work.objects.filter(typing=typing, user_id=user_id, index=index, lesson=lesson).order_by("-id")
        work3 = works[0]
        pic = work3.id

    if request.method == 'POST':
        form = ScoreForm(request.user, request.POST)
        if form.is_valid():
            works = Work.objects.filter(typing=typing, index=index, user_id=user_id, lesson=lesson)
            if works.exists():
                if works[0].score < 0 :
                        # 小老師
                        if not is_teacher(request.user, classroom_id):
                            # credit
                            update_avatar(request.user.id, 2, 1)
                            # History
                            history = PointHistory(user_id=request.user.id, kind=2, message=u'1分--小老師:<'+lesson_name+u'><'+enroll.student.first_name+u'>', url="/student/work/show/"+str(lesson)+"/"+str(index))
                            history.save()

                        # credit
                        update_avatar(enroll.student_id, 1, 1)
                        # History
                        history = PointHistory(user_id=user_id, kind=1, message=u'1分--作業受評<'+lesson_name+u'><'+request.user.first_name+u'>', url="/student/work/show/"+str(lesson)+"/"+str(index))
                        history.save()

                works.update(score=form.cleaned_data['score'])
                works.update(comment=form.cleaned_data['comment'])
                works.update(scorer=request.user.id)
				
                if form.cleaned_data['comment']:
                    # create Message
                    title = u"<" + request.user.first_name+ u">給了評語<" + lesson_name + u">"
                    url = "/student/work/show/" + str(typing) + "/" + str(index) + "/" + str(enroll.student_id)
                    message = Message(title=title, url=url, time=timezone.now())
                    message.save()		

                    # message for group member
                    messagepoll = MessagePoll(message_id = message.id,reader_id=enroll.student_id)
                    messagepoll.save()	

            if is_teacher(request.user, classroom_id):
                if form.cleaned_data['assistant']:
                    try :
                         assistant = WorkAssistant.objects.get(typing=typing, student_id=user_id, group=group, index=index, lesson=lesson, classroom_id=classroom_id)
                    except ObjectDoesNotExist:
                        assistant = WorkAssistant(typing=typing, student_id=user_id, group=group, index=index, lesson=lesson, classroom_id=classroom_id)
                        assistant.save()

                    group_id = WorkGroup.objects.get(classroom_id=classroom_id, index=index, typing=typing).group_id
                    enroll_id = Enroll.objects.get(classroom_id=classroom_id, student_id=user_id).id
                    try :
                        studentgroup = StudentGroup.objects.get(enroll_id=enroll_id, group_id=group_id)
                        group = studentgroup.group
                        studentgroups = StudentGroup.objects.filter(group_id=group_id, group=group)    
                        enroll_ids = map(lambda a: a.enroll_id, studentgroups)                                             
                    except ObjectDoesNotExist:
                        group = -1
                        enroll_ids = [enroll_id]

                    # create Message
                    title = u"<" + assistant.student.first_name+ u">擔任小老師<" + lesson_name + u">"
                    url = "/teacher/score_peer/" + str(typing) + "/" + str(index) + "/" + str(classroom_id) + "/" + str(group)
                    message = Message(title=title, url=url, time=timezone.now())
                    message.save()

                    enrolls = Enroll.objects.filter(id__in=enroll_ids)
                    for enroll in enrolls:
                            # message for group member
                            messagepoll = MessagePoll(message_id = message.id,reader_id=enroll.student_id)
                            messagepoll.save()
                if typing == 0:
                    return redirect('/teacher/work/class/'+str(typing)+ "/"+str(classroom_id)+'/'+str(index))
                elif typing == 1:
                    return redirect('/teacher/work2/class/'+str(classroom_id)+'/'+str(index))  
            else:
                group_id = WorkGroup.objects.get(classroom_id=classroom_id, index=index, typing=typing).group_id
                enroll_id = Enroll.objects.get(classroom_id=classroom_id, student_id=request.user.id).id
                try :
                    studentgroup = StudentGroup.objects.get(enroll_id=enroll_id, group_id=group_id)
                    group = studentgroup.group                                      
                except ObjectDoesNotExist:
                    group = -1                                
                return redirect('/teacher/score_peer/'+str(typing)+"/"+str(index)+'/'+str(classroom_id)+'/'+str(group))

    else:
        works = Work.objects.filter(typing=typing, index=index, user_id=user_id, lesson=lesson).order_by("-id")
        if not works.exists():
            form = ScoreForm(user=request.user)
        else:
            form = ScoreForm(instance=works[0], user=request.user)
            workfiles = WorkFile.objects.filter(work_id=works[0].id).order_by("-id")
        try:
            group_id = WorkGroup.objects.get(typing=typing, classroom_id=classroom_id, index=index).group_id
        except ObjectDoesNotExist:
            group_id = 0
    return render(request, 'teacher/scoring.html', {'typing':typing, 'form': form,'work':work3, 'pic':pic, 'workfiles':workfiles, 'teacher':teacher, 'student':user, 'classroom_id':classroom_id, 'lesson':lesson, 'index':index, 'group_id': group_id})

# 小老師評分名單
def score_peer(request, typing, index, classroom_id, group):
    classroom = Classroom.objects.get(id=classroom_id)
    lesson = classroom.lesson    
    try:
        group_id = WorkGroup.objects.get(typing=typing, classroom_id=classroom_id, index=index).group_id
    except ObjectDoesNotExist:
        group_id = 0    
    if typing == 0:
        lesson_dict = OrderedDict()
        for unit1 in lesson_list[int(lesson)-1][1]:
            for assignment in unit1[1]:
                lesson_dict[assignment[2]] = assignment[0]
        queryset = lesson_dict[int(index)]
    elif typing == 1:
        queryset = TWork.objects.get(id=index).title
    try:
        assistant = WorkAssistant.objects.get(typing=typing, lesson=lesson, index=index, group=group, student_id=request.user.id)
    except ObjectDoesNotExist:
        if typing == 0:
            return redirect("/student/group/work/0/"+str(lesson)+"/"+str(index)+"/"+str(classroom_id))
        elif typing == 1:
            return redirect("/student/group/work/1/"+str(lesson)+"/"+str(index)+"/"+str(classroom_id))
    except MultipleObjectsReturned:
        assistant = WorkAssistant.objects.filter(typing=typing, lesson=lesson, index=index, group=group, student_id=request.user.id).order_by("-id")[0]
    workgroup = WorkGroup.objects.get(classroom_id=classroom_id, index=index, typing=typing)
    studentgroups = StudentGroup.objects.filter(group_id=workgroup.group_id, group=group)
    enroll_ids = map(lambda a: a.enroll_id, studentgroups) 
    enrolls = Enroll.objects.filter(id__in=enroll_ids)
    lessons = ""
    classmate_work = []
    for enroll in enrolls:
        if not enroll.student_id == request.user.id :
            scorer_name = ""
            try:
                work = Work.objects.get(typing=typing, user_id=enroll.student.id, index=index, lesson=lesson)
                if work.scorer > 0 :
                    scorer = User.objects.get(id=work.scorer)
                    scorer_name = scorer.first_name
            except ObjectDoesNotExist:
                work = Work(typing=typing, index=index, user_id=enroll.student.id, lesson=lesson)
            except MultipleObjectsReturned:
                work = Work.objects.filter(typing=typing, user_id=enroll.student.id, index=index, lesson=lesson).order_by("-id")[0]
            classmate_work.append([enroll.student,work,1, scorer_name])
            lessons = queryset
    return render(request, 'teacher/score_peer.html',{'lessons':lessons, 'enrolls':enrolls, 'classmate_work': classmate_work, 'classroom_id':classroom_id, 'lesson':lesson, 'index': index, 'typing':typing})

# 心得
def memo(request, classroom_id):
    if not is_teacher(request.user, classroom_id):
        return redirect("/")
    enrolls = Enroll.objects.filter(classroom_id=classroom_id).order_by("seat")
    classroom = Classroom.objects.get(id=classroom_id)
    return render(request, 'teacher/memo.html', {'lesson':classroom.lesson, 'enrolls':enrolls, 'classroom_name':classroom.name})

# 評分某同學某進度心得
@login_required
def check(request, typing, unit, user_id, classroom_id):
    if not is_teacher(request.user, classroom_id):
        return redirect("/")
    classroom = Classroom.objects.get(id=classroom_id)
    lesson = classroom.lesson
    user_name = User.objects.get(id=user_id).first_name
    lesson_dict = {}
    works = Work.objects.filter(typing=typing, user_id=user_id, lesson=lesson)
    enroll = Enroll.objects.get(student_id=user_id, classroom_id=classroom_id)
    if typing == 0:
        if lesson == 1:
            for assignment in lesson_list1:
                lesson_dict[assignment[3]] = [assignment]
        elif lesson == 2:
            for unit1 in lesson_list[int(lesson)-2][1]:
                for assignment in unit1[1]:
                    lesson_dict[assignment[2]] = [assignment]
    else :
        assignments = TWork.objects.filter(classroom_id=classroom_id)
        for assignment in assignments:
            lesson_dict[assignment.id] = [assignment]

    for work in works:
        index = work.index
        if index in lesson_dict:
            lesson_dict[index].append(work.score)
            lesson_dict[index].append(work.publication_date)
            if work.score > 0:
                score_name = User.objects.get(id=work.scorer).first_name
                lesson_dict[index].append(score_name)
            else :
                lesson_dict[index].append("尚未評分!")
            lesson_dict[index].append(work.memo)
    if request.method == 'POST':
        if typing == 0:        
            form = CheckForm(request.POST)
        else:
            form = CheckForm2(request.POST)
        if form.is_valid():
            if typing == 0:
                enroll.score_memo0 = form.cleaned_data['score_memo0']
                enroll.score_memo1 = form.cleaned_data['score_memo1']
                enroll.score_memo2 = form.cleaned_data['score_memo2']                            
            else :
                enroll.score_memo0_custom = form.cleaned_data['score_memo0_custom']
                enroll.score_memo1_custom = form.cleaned_data['score_memo1_custom']
                enroll.score_memo2_custom = form.cleaned_data['score_memo2_custom']
            enroll.save()            
            return redirect('/teacher/memo/'+str(classroom_id))
    else:
        if typing == 0:
            form = CheckForm(instance=enroll)
        else :
            form = CheckForm2(instance=enroll)
    return render(request, 'teacher/check.html', {'typing':typing, 'works':works, 'lesson': lesson, 'unit':unit, 'form':form, 'works':works, 'lesson_list':sorted(lesson_dict.items()), 'enroll': enroll, 'classroom_id':classroom_id})

# 退選
@login_required
def unenroll(request, enroll_id, classroom_id):
    if not in_teacher_group(request.user):
        return redirect("/")
    enroll = Enroll.objects.get(id=enroll_id)
    enroll.delete()
    classroom_name = Classroom.objects.get(id=classroom_id).name

    return redirect('/student/classroom/'+str(classroom_id)+'/classmate/')
	
# 列出所有公告
class AnnounceListView(ClassroomTeacherRequiredMixin, ListView):
    model = Message
    context_object_name = 'messages'
    template_name = 'teacher/announce_list.html'
    paginate_by = 20

    def get_queryset(self):
        queryset = Message.objects.filter(type=1, classroom_id=self.kwargs['classroom_id'], author_id=self.request.user.id).order_by("-id")
        return queryset

    def get_context_data(self, **kwargs):
        context = super(AnnounceListView, self).get_context_data(**kwargs)
        context['classroom'] = Classroom.objects.get(id=self.kwargs['classroom_id'])
        return context

    # 限本班任課教師
    def render(request, self, context):
        if not is_teacher(self.request.user, self.kwargs['classroom_id']) and not is_assistant(self.request.user, self.kwargs['classroom_id']):
            return redirect('/')
        return super(AnnounceListView, self).render(request, context)

#新增一個公告
class AnnounceCreateView(CreateView):
    model = Message
    form_class = AnnounceForm
    template_name = 'teacher/announce_form.html'  

    def form_valid(self, form):
        self.object = form.save(commit=False)			
        classrooms = self.request.POST.getlist('classrooms')
        files = []
        if self.request.FILES.getlist('files'):
             for file in self.request.FILES.getlist('files'):
                fs = FileSystemStorage()
                filename = uuid4().hex							
                fs.save("static/upload/"+str(self.request.user.id)+"/"+filename, file)								
                files.append([filename, file.name])		
        for classroom_id in classrooms:
            message = Message()
            message.title = u"[公告]" + self.object.title
            message.author_id = self.request.user.id	
            message.type = 1 #公告
            message.classroom_id = classroom_id
            message.content = self.object.content
            message.save()
            message.url = "/account/line/detail/" + classroom_id + "/" + str(message.id)
            message.save()
            if files:
                for file, name in files:
                    content = MessageContent()
                    content.title = name
                    content.message_id = message.id
                    content.filename = str(self.request.user.id)+"/"+file
                    content.save()		

            # 班級學生訊息
            enrolls = Enroll.objects.filter(classroom_id=classroom_id)
            for enroll in enrolls:
                messagepoll = MessagePoll(message_type=1, message_id=message.id, reader_id=enroll.student_id, classroom_id=classroom_id)
                messagepoll.save()               
        return redirect("/student/announce/"+str(self.kwargs['classroom_id'])) 
			
    def get_context_data(self, **kwargs):
        context = super(AnnounceCreateView, self).get_context_data(**kwargs)
        context['class'] = Classroom.objects.get(id=self.kwargs['classroom_id'])
        classroom_list = []
        classrooms = Classroom.objects.filter(teacher_id=self.request.user.id)
        for classroom in classrooms:
            classroom_list.append(classroom.id)
        assistants = Assistant.objects.filter(user_id=self.request.user.id)
        for assistant in assistants:
            if not assistant.classroom_id in classroom_list:
                classroom_list.append(assistant.classroom_id)
        classrooms = Classroom.objects.filter(id__in=classroom_list).order_by("-id")[:30]
        context['classrooms'] = classrooms
        return context	   
        

# 公告
def announce_detail(request, classroom_id, message_id):
    if not is_classmate(request.user, classroom_id):
        return redirect("/")
    message = Message.objects.get(id=message_id)
    files = MessageContent.objects.filter(message_id=message_id)
    classroom = Classroom.objects.get(id=message.classroom_id)

    announce_reads = []

    messagepolls = MessagePoll.objects.filter(message_id=message_id)
    for messagepoll in messagepolls:
        try:
            enroll = Enroll.objects.get(classroom_id=message.classroom_id, student_id=messagepoll.reader_id)
            announce_reads.append([enroll.seat, enroll.student.first_name, messagepoll])
        except ObjectDoesNotExist:
            pass

    def getKey(custom):
        return custom[0]
    announce_reads = sorted(announce_reads, key=getKey)
    return render(request, 'teacher/announce_detail.html', {'files':files,'message':message, 'classroom':classroom, 'announce_reads':announce_reads})

	
# Ajax 設為小教師、取消小教師
def steacher_make(request):
    classroom_id = request.POST.get('classroomid')
    user_id = request.POST.get('userid')
    action = request.POST.get('action')
    lesson = request.POST.get('lesson')
    typing = request.POST.get('typing')	
    index = request.POST.get('index')

    lesson_dict = OrderedDict()
    if typing == "0":
        for unit1 in lesson_list[int(lesson)-1][1]:
            for assignment in unit1[1]:
                lesson_dict[assignment[2]] = assignment[0]
        assignment = lesson_dict[int(index)]
    else :
        assignment = TWork.objects.get(id=index).title
		
    if is_teacher(request.user, classroom_id):
        if user_id and action and lesson and index and typing:
            user = User.objects.get(id=user_id)
        try:
            group_id = WorkGroup.objects.get(typing=typing, classroom_id=classroom_id, index=index).group_id
        except ObjectDoesNotExist:
            group_id = 0            
        group_id = WorkGroup.objects.get(classroom_id=classroom_id, index=index, typing=typing).group_id
        enroll_id = Enroll.objects.get(classroom_id=classroom_id, student_id=user_id).id
        try :
            studentgroup = StudentGroup.objects.get(enroll_id=enroll_id, group_id=group_id)
            group = studentgroup.group
        except ObjectDoesNotExist:
            group = -1  
        if action == "set":
            try:
                assistant = WorkAssistant.objects.get(typing=typing, group=group, student_id=user_id, lesson=lesson, index=index)
            except ObjectDoesNotExist:
                assistant = WorkAssistant(typing=typing, group=group, student_id=user_id, lesson=lesson, index=index)
            assistant.save()


            # create Message
            title = u"<" + assistant.student.first_name+ u">擔任小老師<" + assignment + u">"
            url = "/teacher/score_peer/" + str(typing) + "/" + str(index) + "/" + str(classroom_id) + "/" + str(group)
            message = Message(title=title, url=url, time=timezone.now())
            message.save()

        else:
            try:
                assistant = WorkAssistant.objects.get(typing=typing, student_id=user_id, group=group, lesson=lesson, index=index)
                assistant.delete()					
            except ObjectDoesNotExist:
                assistant = WorkAssistant(student_id=user_id, group=group, lesson=lesson, index=index)
			
            # create Message
            title = "<" + assistant.student.first_name + u">取消小老師<" + assignment + ">"
            url = "/student/group/work/" + str(typing) + "/" + str(lesson) + "/" + str(index) + "/" + str(classroom_id)
            message = Message(title=title, url=url, time=timezone.now())
            message.save()
        if group >= 0 :
            studentgroups = StudentGroup.objects.filter(group_id=group_id, group=group)    
            enroll_ids = map(lambda a: a.enroll_id, studentgroups)                                             
            enrolls = Enroll.objects.filter(id__in=enroll_ids)
            for enroll in enrolls:
                # message for group member
                messagepoll = MessagePoll(message_id = message.id,reader_id=enroll.student_id)
                messagepoll.save()
        return JsonResponse({'status':'ok'}, safe=False)
    else:
        return JsonResponse({'status':classroom_id}, safe=False)	
		
def grade(request, typing, unit, classroom_id):
    # 限本班任課教師
    if not is_teacher(request.user, classroom_id):
        return redirect("/")
    enrolls = Enroll.objects.filter(classroom_id=classroom_id).order_by('seat')
    classroom = Classroom.objects.get(id=classroom_id)     
    lesson = classroom.lesson
    user_ids = [enroll.student_id for enroll in enrolls]
    work_pool = Work.objects.filter(typing=typing, user_id__in=user_ids, lesson=lesson).order_by('id')
    lesson_dict = {}
    data = []
	
    lesson_dict = OrderedDict()
    for unit1 in lesson_list[int(lesson)-1][1]:
        for assignment in unit1[1]:
            lesson_dict[assignment[2]] = assignment
    lesson_list2 = sorted(lesson_dict.items())
    for enroll in enrolls:
      enroll_score = []
      total = 0
      memo = 0 
      grade = 0
      if typing == 0:
        lesson_list2 = lesson_list2
      elif typing == 1:
        lesson_list2 = TWork.objects.filter(classroom_id=classroom_id)
      for index, assignment in enumerate(lesson_list2):
            if typing == 0: 			    
                works = list(filter(lambda w: w.index == index+1 and w.user_id==enroll.student_id, work_pool))
            else :
                works = list(filter(lambda w: w.index == assignment.id and w.user_id==enroll.student_id, work_pool))
            works_count = len(works)
            if works_count == 0:
                enroll_score.append(["X", index])
                if typing == 0 or typing == 1:
                    total += 60
            else:
                work = works[-1]
                enroll_score.append([work.score, index])
                if work.score == -1:
                    if typing == 0 or typing == 1:
                        total += 80
                else:
                    total += work.score

            if typing == 0:
                memo0 = enroll.score_memo0
                memo1 = enroll.score_memo1
                memo2 = enroll.score_memo2                                
            elif typing == 1:
                memo0 = enroll.score_memo0_custom
                memo1 = enroll.score_memo1_custom
                memo2 = enroll.score_memo2_custom
            memo = [memo0, memo1, memo2]                                
            grade = int(total / len(lesson_list2) * 0.6 + memo[0] * 0.4)
      data.append([enroll, enroll_score, memo, grade])
    return render(request, 'teacher/grade.html', {'typing':typing, 'unit':unit, 'lesson_list':lesson_list2, 'classroom':classroom, 'data':data})

def grade_excel(request, typing, unit, classroom_id):
    # 限本班任課教師
    if not is_teacher(request.user, classroom_id) and not is_assistant(request.user, classroom_id):
        return redirect("/")
    enrolls = Enroll.objects.filter(classroom_id=classroom_id).order_by('seat')
    classroom = Classroom.objects.get(id=classroom_id)     
    lesson = classroom.lesson    
    user_ids = [enroll.student_id for enroll in enrolls]
    work_pool = Work.objects.filter(typing=typing, user_id__in=user_ids, lesson=lesson).order_by('id')
    lesson_dict = OrderedDict()
    data = []
    for unit1 in lesson_list[int(lesson)-1][1]:
        for assignment in unit1[1]:
            lesson_dict[assignment[2]] = assignment
    lesson_list2 = sorted(lesson_dict.items())
    for enroll in enrolls:
      enroll_score = []
      total = 0
      memo = 0 
      grade = 0
      stu_works = filter(lambda w: w.user_id == enroll.student_id, work_pool)
      if typing == 0:
        lesson_list2 = lesson_list2
      elif typing == 1:
        lesson_list2 = TWork.objects.filter(classroom_id=classroom_id)
      for index, assignment in enumerate(lesson_list2):
            if typing == 0: 			    
                works = list(filter(lambda w: w.index == index+1, stu_works))
            else :
                works = list(filter(lambda w: w.index == assignment.id, stu_works))
            works_count = len(works)
            if works_count == 0:
                enroll_score.append(["X", index])
                if typing == 0 or typing == 1:
                    total += 60
            else:
                work = works[-1]
                enroll_score.append([work.score, index])
                if work.score == -1:
                    if typing == 0 or typing == 1:
                        total += 80
                else:
                    total += work.score

            if typing == 0:
                memo0 = enroll.score_memo0
                memo1 = enroll.score_memo1
                memo2 = enroll.score_memo2                                
            elif typing == 1:
                memo0 = enroll.score_memo0_custom
                memo1 = enroll.score_memo1_custom
                memo2 = enroll.score_memo2_custom
            memo = [memo0, memo1, memo2] 
            grade = int(total / len(lesson_list2) * 0.6 + memo[0] * 0.4)
      data.append([enroll, enroll_score, memo, grade])
                
    output = io.BytesIO()
    workbook = xlsxwriter.Workbook(output)    
    worksheet = workbook.add_worksheet(classroom.name)
    date_format = workbook.add_format({'num_format': 'yy/mm/dd'})

    row = 1
    worksheet.write(row, 1, u'座號')
    worksheet.write(row, 2, u'姓名')
    worksheet.write(row, 3, u'成績')        
    worksheet.write(row, 4, u"心得(全)")
    worksheet.write(row, 5, u"心得(上)") 
    worksheet.write(row, 6, u"心得(下)")       
    index = 7
    for assignment in lesson_list2:
        if typing == 0:
	          worksheet.write(row, index, assignment[1][0])
        else :
	          worksheet.write(row, index, assignment.title)              
        index += 1


    index = 4
    if not typing == 0:
        row += 1
        for assignment in lesson_list2:            
            worksheet.write(row, index, datetime.strptime(str(assignment.time)[:19],'%Y-%m-%d %H:%M:%S'), date_format)
            index += 1			

    for enroll, enroll_score, memo, grade in data:
      row += 1
      worksheet.write(row, 1, enroll.seat)
      worksheet.write(row, 2, enroll.student.first_name)
      worksheet.write(row, 3, grade)     
      worksheet.write(row, 4, memo[0])
      worksheet.write(row, 5, memo[1])
      worksheet.write(row, 6, memo[2])            
      index = 7
      for score, index2 in enroll_score:
          if score == -1 :
              worksheet.write(row, index, "V")
          else :
              worksheet.write(row, index, score)
          index +=1 
 

    workbook.close()
    # xlsx_data contains the Excel file
    response = HttpResponse(content_type='application/vnd.ms-excel')
    if typing == 0:
        type_name = "指定作業"
    elif typing == 1:
        type_name = "自訂作業"      
    filename = classroom.name + '-' + type_name + "-" + str(localtime(timezone.now()).date()) + '.xlsx'
    response['Content-Disposition'] = filename_browser(request, filename)
    xlsx_data = output.getvalue()
    response.write(xlsx_data)
    return response		
	
# 列出班級各作業小老師
def work_assistant(request, typing, classroom_id):
    # 限本班任課教師
    if not is_teacher(request.user, classroom_id):
        return redirect("/")    
    classroom = Classroom.objects.get(id=classroom_id)
    lesson = Classroom.objects.get(id=classroom_id).lesson		
    enroll_pool = [enroll for enroll in Enroll.objects.filter(classroom_id=classroom_id).order_by('seat')]
    enroll_ids =  map(lambda a: a.id, enroll_pool)
    student_ids = map(lambda a: a.student_id, enroll_pool)
    studentgroup_pool = [studentgroup for studentgroup in StudentGroup.objects.filter(enroll_id__in=enroll_ids)]
    work_pool = Work.objects.filter(user_id__in=student_ids, lesson=classroom.lesson, typing=typing)
    user_pool = [user for user in User.objects.filter(id__in=work_pool.values('scorer'))]
    assistant_pool = [assistant for assistant in WorkAssistant.objects.filter(typing=typing, lesson=lesson, classroom_id=classroom_id)]				
    lessons = []		
	
    if typing == 0: 
        lesson_dict = OrderedDict()
        lesson = classroom.lesson
        if lesson == 1:
            assignments = lesson_list1        	
            for assignment in assignments:
                student_groups = []	
                try:
                    workgroup = WorkGroup.objects.get(classroom_id=classroom_id, index=assignment[3], typing=typing)
                    group_id = workgroup.group_id                    
                except ObjectDoesNotExist:
                    group_id = 0
                if group_id > 0:
                    classroomgroup = ClassroomGroup.objects.get(id=workgroup.group_id)
                    groups = range(classroomgroup.numbers)				 
                    for group in groups:
                        studentgroups = filter(lambda u: u.group == group and u.group_id==workgroup.group_id, studentgroup_pool)
                        members = []
                        for studentgroup in studentgroups:
                            member = list(filter(lambda u: u.id == studentgroup.enroll_id, enroll_pool))[0]
                            members.append(member)
                        group_assistants = []
                        works = []
                        scorer_name = ""
                        for member in members:
                            work = list(filter(lambda w: w.index == assignment[2] and w.user_id == member.student_id, work_pool))
                            if work:
                                work = work[0]
                                scorer = list(filter(lambda u: u.id == work.scorer, user_pool))
                                scorer_name = scorer[0].first_name if scorer else 'X'
                            else:
                                work = Work(index=assignment[2], user_id=1, score=-2)
                            works.append([member, work.score, scorer_name, work.memo])
                            assistant = list(filter(lambda a: a.student_id == member.student_id and a.index == assignment[2], assistant_pool))
                            if assistant:
                                group_assistants.append(member)                    
                        group_name = "第"+ str(group+1) + "組"
                        student_groups.append([group, works, group_assistants, group_name])
                lesson_dict[assignment[3]] = [assignment, student_groups, group_id]		
        elif lesson == 2:
            for unit1 in lesson_list[int(lesson)-2][1]:
                for assignment in unit1[1]:
                    student_groups = []	
                    try:
                        workgroup = WorkGroup.objects.get(classroom_id=classroom_id, index=assignment[2], typing=typing)
                        group_id = workgroup.group_id                    
                    except ObjectDoesNotExist:
                        group_id = 0
                    if group_id > 0:
                        classroomgroup = ClassroomGroup.objects.get(id=workgroup.group_id)
                        groups = range(classroomgroup.numbers)				 
                        for group in groups:
                            studentgroups = filter(lambda u: u.group == group and u.group_id==workgroup.group_id, studentgroup_pool)
                            members = []
                            for studentgroup in studentgroups:
                                member = list(filter(lambda u: u.id == studentgroup.enroll_id, enroll_pool))[0]
                                members.append(member)
                            group_assistants = []
                            works = []
                            scorer_name = ""
                            for member in members:
                                work = list(filter(lambda w: w.index == assignment[2] and w.user_id == member.student_id, work_pool))
                                if work:
                                    work = work[0]
                                    scorer = list(filter(lambda u: u.id == work.scorer, user_pool))
                                    scorer_name = scorer[0].first_name if scorer else 'X'
                                else:
                                    work = Work(index=assignment[2], user_id=1, score=-2)
                                works.append([member, work.score, scorer_name, work.memo])
                                assistant = list(filter(lambda a: a.student_id == member.student_id and a.index == assignment[2], assistant_pool))
                                if assistant:
                                    group_assistants.append(member)                    
                            group_name = "第"+ str(group+1) + "組"
                            student_groups.append([group, works, group_assistants, group_name])
                        lesson_dict[assignment[2]] = [assignment, student_groups, group_id]		
                    else:
                        lesson_dict[assignment[2]] = [assignment, [], group_id]
    else :
        lesson_dict = OrderedDict()
        assignments = TWork.objects.filter(classroom_id=classroom_id)
        for assignment in assignments:          
            student_groups = []	
            workgroup = WorkGroup.objects.get(classroom_id=classroom_id, index=assignment.id, typing=typing)
            group_id = workgroup.group_id
            if group_id > 0:
                classroomgroup = ClassroomGroup.objects.get(id=workgroup.group_id)
                groups = range(classroomgroup.numbers)            
                student_groups = []													
                for group in groups:
                    studentgroups = filter(lambda u: u.group == group and u.group_id==workgroup.group_id, studentgroup_pool)
                    members = []
                    for studentgroup in studentgroups:
                        member = list(filter(lambda u: u.id == studentgroup.enroll_id, enroll_pool))[0]
                        members.append(member)                    
                    members = filter(lambda u: u.group == group, enroll_pool)
                    group_assistants = []
                    works = []
                    scorer_name = ""
                    for member in members:
                        work = list(filter(lambda w: w.index == assignment.id and w.user_id == member.student_id, work_pool))
                        if work:
                            work = work[0]
                            scorer = list(filter(lambda u: u.id == work.scorer, user_pool))
                            scorer_name = scorer[0].first_name if scorer else 'X'
                        else:
                            work = Work(index=assignment.id, user_id=1, score=0)
                        works.append([member, work.score, scorer_name, work.memo])
                        assistant = list(filter(lambda a: a.user_id == member.student_id and a.lesson == assignment.id, assistant_pool))
                        if assistant:
                            group_assistants.append(member)                    
                    group_name = "第"+ str(group+1) + "組"
                    student_groups.append([group, works, group_assistants, group_name])
                lesson_dict[assignment.id] = [assignment, student_groups, group_id]
            else :
                lesson_dict[assignment.id] = [assignment, [], group_id]		
    return render(request, 'teacher/work_groups.html', {'typing':typing, 'test':lesson_dict, 'lesson_dict':sorted(lesson_dict.items()),'classroom':classroom})
	
def survey(request, classroom_id):
    return render(request, 'teacher/survey.html', {'classroom_id':classroom_id})
	
# 列出所有測驗主題
class ExamListView(ClassroomTeacherRequiredMixin, ListView):
    model = Exam
    context_object_name = 'exams'
    template_name = "teacher/exam_list.html"		
    paginate_by = 20
    def get_queryset(self):        
        exam_classes = ExamClass.objects.filter(classroom_id=self.kwargs['classroom_id']).order_by("-publication_date", "-exam_id")
        exams = []
        for exam_class in exam_classes:
            exam = Exam.objects.get(id=exam_class.exam_id)
            exams.append([exam, exam_class])
        return exams
			
    def get_context_data(self, **kwargs):
        context = super(ExamListView, self).get_context_data(**kwargs)
        classroom = Classroom.objects.get(id=self.kwargs['classroom_id'])
        context['classroom'] = classroom
        return context	
        
#新增一個測驗主題
class ExamCreateView(ClassroomTeacherRequiredMixin, CreateView):
    model = Exam
    form_class = ExamForm
    template_name = "teacher/exam_form.html"
    def form_valid(self, form):
        self.object = form.save(commit=False)
        self.object.user_id = self.request.user.id
        self.object.classroom_id = self.kwargs['classroom_id']        
        self.object.save()  
        classrooms = self.request.POST.getlist('classrooms')
        for classroom in classrooms:
          exam_class = ExamClass(exam_id=self.object.id, classroom_id=classroom)
          exam_class.save()
        
        return redirect("/teacher/exam/"+str(self.kwargs['classroom_id']))           
        
    def get_context_data(self, **kwargs):
        context = super(ExamCreateView, self).get_context_data(**kwargs)
        classroom_list = []
        classrooms = Classroom.objects.filter(teacher_id=self.request.user.id)
        for classroom in classrooms:
            classroom_list.append(classroom.id)
        assistants = Assistant.objects.filter(user_id=self.request.user.id)
        for assistant in assistants:
            if not assistant.classroom_id in classroom_list:
                classroom_list.append(assistant.classroom_id)
        classrooms = Classroom.objects.filter(id__in=classroom_list).order_by("-id")
        context['classrooms'] = classrooms
        context['classroom_id'] = int(self.kwargs['classroom_id'])
        context['classroom'] = Classroom.objects.get(id=self.kwargs['classroom_id'])
        return context	
  
        return redirect("/teacher/exam/"+self.kwargs['classroom_id']) 			
	
# 列出某測驗主題的班級
class ExamClassListView(TeacherRequiredMixin, ListView):
    model = Exam
    context_object_name = 'classrooms'
    template_name = "teacher/exam_class.html"		
    paginate_by = 20
	
    def get_queryset(self):        		
        eclass_dict = dict(((eclass.classroom_id, eclass) for eclass in ExamClass.objects.filter(exam_id=self.kwargs['exam_id'])))		
        classroom_list = []
        classroom_ids = []
        classrooms = Classroom.objects.filter(teacher_id=self.request.user.id).order_by("-id")
        for classroom in classrooms:
            if classroom.id in eclass_dict:
                classroom_list.append([classroom, True, eclass_dict[classroom.id].deadline, eclass_dict[classroom.id].deadline_date])
            else :
                classroom_list.append([classroom, False, False, timezone.now()])
            classroom_ids.append(classroom.id)
        assistants = Assistant.objects.filter(user_id=self.request.user.id)
        for assistant in assistants:
            classroom = Classroom.objects.get(id=assistant.classroom_id)
            if not classroom.id in classroom_ids:
                if classroom.id in eclass_dict:
                    classroom_list.append([classroom, True, eclass_dict[classroom.id].deadline, eclass_dict[classroom.id].deadline_date])
                else :
                    classroom_list.append([classroom, False, False, timezone.now()])
        return classroom_list
			
    def get_context_data(self, **kwargs):
        context = super(ExamClassListView, self).get_context_data(**kwargs)				
        exam = Exam.objects.get(id=self.kwargs['exam_id'])
        context['exam'] = exam
        context['exam_id'] = self.kwargs['exam_id']
        return context	
	
# Ajax 開放班取、關閉班級
def exam_switch(request):
    exam_id = request.POST.get('examid')
    classroom_id = request.POST.get('classroomid')		
    status = request.POST.get('status')
    try:
        examclass = ExamClass.objects.get(exam_id=exam_id, classroom_id=classroom_id)
        if status == 'false' :
            examclass.delete()
    except ObjectDoesNotExist:
        if status == 'true':
            examclass = ExamClass(exam_id=exam_id, classroom_id=classroom_id)
            examclass.save()
    return JsonResponse({'status':status}, safe=False)        
	
class ExamEditUpdateView(ClassroomTeacherRequiredMixin, UpdateView):
    model = Exam
    fields = ['title']
    template_name = 'form.html'

    def get_success_url(self):
        succ_url =  '/teacher/exam/'+str(self.kwargs['classroom_id'])
        return succ_url
		
def exam_deadline(request, classroom_id, exam_id):
    if not is_teacher(request.user, classroom_id):
        return redirect("/")
    exam = Exam.objects.get(id=exam_id)
    if request.method == 'POST':
        form = ExamDeadlineForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('/teacher/exam/'+str(classroom_id))
    else:
        examclass = ExamClass.objects.get(classroom_id=classroom_id, exam_id=exam_id)
        form = ExamDeadlineForm(instance=examclass)
    return render(request,'teacher/exam_deadline_form.html',{'examclass':examclass, 'exam':exam})

	
# Ajax 設定期限、取消期限
def exam_deadline_set(request):
    exam_id = request.POST.get('examid')
    classroom_id = request.POST.get('classroomid')		
    status = request.POST.get('status')
    try:
        examclass = ExamClass.objects.get(exam_id=exam_id, classroom_id=classroom_id)
    except ObjectDoesNotExist:
        examclass = Examclass(exam_id=exam_id, classroom_id=classroom_id)
    if status == 'True':
        examclass.deadline = True
    else :
        examclass.deadline = False
    examclass.save()
    return JsonResponse({'status':status}, safe=False)        

# Ajax 設定期限日期
def exam_deadline_date(request):
    exam_id = request.POST.get('examid')
    classroom_id = request.POST.get('classroomid')		
    deadline_date = request.POST.get('deadlinedate')
    try:
        examclass = ExamClass.objects.get(exam_id=exam_id, classroom_id=classroom_id)
    except ObjectDoesNotExist:
        examclass = ExamClass(exam_id=exam_id, classroom_id=classroom_id)
    #fclass.deadline_date = deadline_date.strftime('%d/%m/%Y')
    examclass.deadline_date = datetime.strptime(deadline_date, '%Y %B %d - %I:%M %p')
    examclass.save()
    return JsonResponse({'status':deadline_date}, safe=False)             
		
# 列出所有測驗題目
class ExamQuestionListView(TeacherRequiredMixin, ListView):
    model = ExamQuestion
    context_object_name = 'questions'
    template_name = "teacher/exam_question.html"		
    def get_queryset(self):
        queryset = ExamQuestion.objects.filter(exam_id=self.kwargs['exam_id']).order_by("-id")
        return queryset
			
    def get_context_data(self, **kwargs):
        context = super(ExamQuestionListView, self).get_context_data(**kwargs)
        exam = Exam.objects.get(id=self.kwargs['exam_id'])
        context['exam']= exam
        context['exam_id'] = self.kwargs['exam_id']
        questions = ExamQuestion.objects.filter(exam_id=self.kwargs['exam_id'])
        context['score_total'] = sum(question.score for question in questions)			
        return context	
			
#新增一個題目
class ExamQuestionCreateView(TeacherRequiredMixin, CreateView):
    model = ExamQuestion
    form_class = ExamQuestionForm
    template_name = "teacher/exam_question_form.html"
    def form_valid(self, form):
        self.object = form.save(commit=False)
        question = ExamQuestion(exam_id=self.object.exam_id)
				#是非題
        if self.object.types == 1:
            question.types = 1
            question.title = self.object.title
            question.answer = self.object.answer
			  #選擇題
        if self.object.types  == 2:
            question.types = 2					
            question.title = self.object.title
            question.option1 = self.object.option1
            question.option2 = self.object.option2
            question.option3 = self.object.option3
            question.option4 = self.object.option4						
            question.answer = self.object.answer
            question.score = self.object.score
        question.save()         
  
        return redirect("/teacher/exam/question/"+self.kwargs['exam_id'])  

    def get_context_data(self, **kwargs):
        ctx = super(ExamQuestionCreateView, self).get_context_data(**kwargs)
        ctx['exam'] = Exam.objects.get(id=self.kwargs['exam_id'])
        return ctx

def exam_question_delete(request, exam_id, question_id):
    if not is_teacher(request.user, classroom_id):
        return redirect("/")
    instance = ExamQuestion.objects.get(id=question_id)
    instance.delete()

    return redirect("/teacher/exam/question/"+str(exam_id))  
	
def exam_question_edit(request, exam_id, question_id):
    if not is_teacher(request.user, classroom_id):
        return redirect("/")
    exam = Exam.objects.get(id=exam_id)
    try:
        instance = ExamQuestion.objects.get(id=question_id)
    except:
        pass
    if request.method == 'POST':
            question_id = request.POST.get("question_id", "")
            try:
                question = ExamQuestion.objects.get(id=question_id)
            except ObjectDoesNotExist:
	              question = ExamQuestion(exam_id= request.POST.get("exam_id", ""), types=form.cleaned_data['types'])
            if question.types == 1:
                question.answer = request.POST.get("answer", "")	
            elif question.types == 2:
                question.option1 = request.POST.get("option1", "")	
                question.option2 = request.POST.get("option2", "")	
                question.option3 = request.POST.get("option3", "")	
                question.option4 = request.POST.get("option4", "")	
                question.score = request.POST.get("score", "")	
                question.answer = request.POST.get("answer", "")	
            question.title = request.POST.get("title", "")
            question.save()
            return redirect('/teacher/exam/question/'+exam_id+"#"+str(question.id))   
    return render(request,'teacher/exam_question_edit.html',{'question': instance, 'exam':exam, 'quesiton_id':question_id})		
			
# Create your views here.
def exam_import_sheet(request, exam_id):
    if not in_teacher_group(request.user):
        return redirect("/")
    #if request.user.id != 1:
    #    return redirect("/")
    if request.method == "POST":
        form = UploadFileForm(request.POST,
                              request.FILES)
        if form.is_valid():
            ExamImportQuestion2.objects.all().delete()
            request.FILES['file'].save_to_database(
                name_columns_by_row=0,
                model=ExamImportQuestion2,
                mapdict=['title', 'option1', 'option2','option3','option4','answer', 'score'])
            questions = ExamImportQuestion2.objects.all()
            return render(request, 'teacher/exam_import_question2.html',{'questions':questions, 'exam_id': exam_id})
        else:
            return HttpResponseBadRequest()
    else:	
        form = UploadFileForm()
    return render(
        request,
        'teacher/exam_upload_form.html',
        {
					  'exam_id': exam_id, 
            'form': form,
            'title': 'Excel file upload and download example',
            'header': ('Please choose any excel file ' +
                       'from your cloned repository:')
        })
	
# Create your views here.
def exam_import_question(request, exam_id):
    if not in_teacher_group(request.user):
        return redirect("/")
           
    questions = ExamImportQuestion2.objects.all()
    for question in questions:
            new_question = ExamQuestion(exam_id=exam_id, types=2, title=question.title, option1=question.option1, option2=question.option2, option3=question.option3, option4=question.option4, answer=question.answer, score=question.score)
            new_question.save()
            
    return redirect('/teacher/exam/question/'+str(exam_id))			
	
def exam_round(request, classroom_id, exam_id):
    if not is_teacher(request.user, classroom_id):
        return redirect("/")
    if not is_teacher(request.user, classroom_id):
        return redirect("/")
    examclass = ExamClass.objects.get(classroom_id=classroom_id, exam_id=exam_id)
    return render(request,'teacher/exam_round.html',{'examclass':examclass})		
	
def exam_round_set(request):
    exam_id = request.POST.get('examid')
    classroom_id = request.POST.get('classroomid')		
    round_limit = request.POST.get('round_limit')
    try:
        examclass = ExamClass.objects.get(exam_id=exam_id, classroom_id=classroom_id)
    except ObjectDoesNotExist:
        examclass = Examclass(exam_id=exam_id, classroom_id=classroom_id)
    examclass.round_limit = int(round_limit)
    examclass.save()
    return JsonResponse({'status':'ok'}, safe=False)  	
	
def exam_score(request, classroom_id, exam_id):
    if not is_teacher(request.user, classroom_id):
        return redirect("/")
    exam = Exam.objects.get(id=exam_id)
    classroom = Classroom.objects.get(id=classroom_id)
    examclass = ExamClass.objects.get(classroom_id=classroom_id, exam_id=exam_id)
    enrolls = Enroll.objects.filter(classroom_id=classroom_id).order_by("seat")
    enroll_ids = []
    for enroll in enrolls:
        enroll_ids.append(enroll.student_id)
    examworks = ExamWork.objects.filter(exam_id=exam_id, student_id__in=enroll_ids, publish=True).order_by("-id")
    scores = []
    for enroll in enrolls:
        works = list(filter(lambda w: w.student_id == enroll.student_id, examworks))
        if len(works) > 0 :
            score_max = max(work.score for work in works)
            score_avg = sum(work.score for work in works) / len(works)	
        else :
            score_max = 0
            score_avg = 0
        scores.append([enroll, works, score_avg, score_max])
    return render(request,'teacher/exam_score.html',{'classroom': classroom, 'exam':exam, 'scores':scores})			

# 列出所有組別
class GroupListView(ListView):
    model = Group
    context_object_name = 'groups'
    template_name = 'teacher/group.html'
    paginate_by = 25
    def get_queryset(self):      
        queryset = ClassroomGroup.objects.filter(classroom_id=self.kwargs['classroom_id']).order_by("-id")
        return queryset
			
    def get_context_data(self, **kwargs):
        context = super(GroupListView, self).get_context_data(**kwargs)
        context['classroom_id'] = self.kwargs['classroom_id']
        return context				
			
#新增一個分組
class GroupCreateView(CreateView):
    model = ClassroomGroup
    form_class = GroupForm
    template_name = 'teacher/group_form.html'    
    def form_valid(self, form):
        self.object = form.save(commit=False)
        self.object.classroom_id = self.kwargs['classroom_id']
        if is_teacher(self.request.user, self.kwargs['classroom_id']) or is_assistant(self.request.user, self.kwargs['classroom_id']):
            self.object.save()
            # 隨機分組
            if form.cleaned_data['assign'] == 1:
                enrolls = Enroll.objects.filter(classroom_id=self.kwargs['classroom_id'], seat__gt=0).order_by('?')
                number = 0
                for enroll in enrolls:
                    group = StudentGroup(group_id=self.object.id, enroll_id=enroll.id, group=(number % self.object.numbers))
                    group.save()
                    number += 1
                self.object.opening=False
                self.object.save()
                
        return redirect("/student/group/list/"+ str(self.object.id))   
			
    def get_context_data(self, **kwargs):
        context = super(GroupCreateView, self).get_context_data(**kwargs)
        return context	
			
class GroupUpdateView(UpdateView):
    model = ClassroomGroup
    form_class = GroupForm2		
    template_name = 'form.html'
    def get_success_url(self):
        succ_url =  '/student/group/list/'+str(self.kwargs['pk'])
        return succ_url
			
    def form_valid(self, form):
        if is_teacher(self.request.user, self.kwargs['classroom_id']) or is_assistant(self.request.user, self.kwargs['classroom_id']):
            group = ClassroomGroup.objects.get(id=self.kwargs['pk'])
            reduce = group.numbers - form.cleaned_data['numbers']
            if reduce > 0:
                for i in range(reduce):
                    StudentGroup.objects.filter(group_id=self.kwargs['pk'], group=group.numbers-i).delete()
            form.save()
        return HttpResponseRedirect(self.get_success_url())
			

# 分組
def make(request):
    group_id = request.POST.get('groupid')
    action = request.POST.get('action')
    if group_id and action :      
        group = ClassroomGroup.objects.get(id=group_id)	
        if is_teacher(request.user, group.classroom_id) or is_assistant(request.user, group.classroom_id):
            if action == "1":            
                group.opening = True   
            else : 
                group.opening = False
            group.save()      
        return JsonResponse({'status':'ok'}, safe=False)
    else:
        return JsonResponse({'status':'fail'}, safe=False) 
			
# 分組
def make2(request, group_id, action):
        group = ClassroomGroup.objects.get(id=group_id)	  
        if is_teacher(request.user, group.classroom_id) or is_assistant(request.user, group.classroom_id):
            if action == 1:            
                group.opening = True
            else : 
                group.opening = False
            group.save()    
        return redirect("/student/group/list/"+str(group.id))

def group_assign(request, classroom_id, group_id):
    if not is_teacher(request.user, classroom_id):
        return redirect("/")
    group = ClassroomGroup.objects.get(id=group_id)
    GroupModelFormset = modelformset_factory(Enroll, fields=['id', 'seat','student_id', 'group'],extra=0)	
    if request.method == 'POST':
        formset = GroupModelFormset(request.POST)
        if formset.is_valid():
            for form in formset:
                try:
                    student = StudentGroup.objects.get(enroll_id=form.cleaned_data["id"].id, group_id=group_id)
                except ObjectDoesNotExist:
                    student = StudentGroup(enroll_id=form.cleaned_data["id"].id, group_id=group_id)
                student.group = form.cleaned_data["group"]            
                student.save()
                form.save()
            return redirect('/student/group/list/'+str(group_id))
        else:
            return redirect("/")
    else:
        enrolls = Enroll.objects.filter(classroom_id=classroom_id)
        for enroll in enrolls:
            try:
                student = StudentGroup.objects.get(enroll_id=enroll.id, group_id=group_id)
                enroll.group = student.group
            except ObjectDoesNotExist:
                enroll.group = -1
            except MultipleObjectsReturned:
                student = StudentGroup.objects.filter(enroll_id=enroll.id, group_id=group_id)[0]
                enroll.group = student.group
            enroll.save()
        formset = GroupModelFormset(queryset=Enroll.objects.filter(classroom_id=classroom_id).order_by("seat"))		
    return render(request, 'teacher/group_assign.html', {'formset': formset, 'group_numbers':range(group.numbers)})		
		
'''    
----------------------- 思辨區
'''
# 列出所有思辨主題
class SpeculationListView(ListView):
    model = SpeculationWork
    context_object_name = 'forums'
    template_name = "teacher/speculation_list.html"		
    paginate_by = 20
    def get_queryset(self):        
        fclasses = SpeculationClass.objects.filter(classroom_id=self.kwargs['classroom_id']).order_by("-publication_date", "-forum_id")
        forums = []
        for fclass in fclasses:
            forum = SpeculationWork.objects.get(id=fclass.forum_id)
            forums.append([forum, fclass])
        return forums
			
    def get_context_data(self, **kwargs):
        context = super(SpeculationListView, self).get_context_data(**kwargs)
        classroom = Classroom.objects.get(id=self.kwargs['classroom_id'])
        context['classroom'] = classroom
        return context	
        
#新增一個討論主題
class SpeculationCreateView(CreateView):
    model = SpeculationWork
    form_class = SpeculationForm
    template_name = "teacher/speculation_form.html"
    def form_valid(self, form):
        self.object = form.save(commit=False)
        self.object.teacher_id = self.request.user.id
        self.object.classroom_id = self.kwargs['classroom_id']        
        self.object.save()  
        classrooms = self.request.POST.getlist('classrooms')
        for classroom in classrooms:
          forum_class = SpeculationClass(forum_id=self.object.id, classroom_id=classroom)
          forum_class.save()
        
        return redirect("/teacher/speculation/"+str(self.kwargs['classroom_id']))           
        
    def get_context_data(self, **kwargs):
        context = super(SpeculationCreateView, self).get_context_data(**kwargs)
        classroom_list = []
        classrooms = Classroom.objects.filter(teacher_id=self.request.user.id)
        for classroom in classrooms:
            classroom_list.append(classroom.id)
        assistants = Assistant.objects.filter(user_id=self.request.user.id)
        for assistant in assistants:
            if not assistant.classroom_id in classroom_list:
                classroom_list.append(assistant.classroom_id)
        classrooms = Classroom.objects.filter(id__in=classroom_list).order_by("-id")
        context['classrooms'] = classrooms
        context['classroom_id'] = int(self.kwargs['classroom_id'])
        context['classroom'] = Classroom.objects.get(id=self.kwargs['classroom_id'])
        return context	
  
        return redirect("/teacher/speculation/"+self.kwargs['classroom_id'])        
	
	
# 列出所有思辨主題
class SpeculationAllListView(ListView):
    model = SpeculationWork
    context_object_name = 'forums'
    template_name = "teacher/speculation_all.html"		
    paginate_by = 20
		
    def get_queryset(self):
      # 年級
      if self.kwargs['categroy'] == "1":
        queryset = SpeculationWork.objects.filter(levels__contains=self.kwargs['categroy_id']).order_by("-id")
      # 學習領域
      elif self.kwargs['categroy'] == "2":
        queryset = SpeculationWork.objects.filter(domains__contains=self.kwargs['categroy_id']).order_by("-id")   
      else:
        queryset = SpeculationWork.objects.all().order_by("-id")
      if self.request.GET.get('account') != None:
        keyword = self.request.GET.get('account')
        users = User.objects.filter(Q(username__icontains=keyword) | Q(first_name__icontains=keyword)).order_by('-id')
        user_list = []
        for user in users:
            user_list.append(user.id)
        forums = queryset.filter(teacher_id__in=user_list)
        return forums
      else:				
        return queryset
			
    def get_context_data(self, **kwargs):
        context = super(SpeculationAllListView, self).get_context_data(**kwargs)
        context['categroy'] = self.kwargs['categroy']							
        context['categroy_id'] = self.kwargs['categroy_id']							
        return context	

# 展示思辨素材
def speculation_show(request, forum_id):
    forum = SpeculationWork.objects.get(id=forum_id)
    domains = Domain.objects.all()
    domain_dict = {}
    for domain in domains :
        key = domain.id
        domain_dict[key] = domain
    levels = Level.objects.all()	
    level_dict = {}
    for level in levels :
        key = level.id
        level_dict[key] = level
    contents = SpeculationContent.objects.filter(forum_id=forum_id)
    domains = []		
    if forum.domains:
        forum_domains = ast.literal_eval(forum.domains)
        for domain in forum_domains:
            key = int(domain)
            domains.append(domain_dict[key])
    levels = []						
    if forum.levels:
        forum_levels = ast.literal_eval(forum.levels)
        for level in forum_levels:
            key = int(level)			
            levels.append(level_dict[key])
    return render(request,'teacher/speculation_show.html',{'domains':domains, 'levels':levels, 'contents':contents, 'forum':forum})

		
# 列出某思辨主題的班級
class SpeculationClassListView(ListView):
    model = SpeculationWork
    context_object_name = 'classrooms'
    template_name = "teacher/speculation_class.html"		
    paginate_by = 20
	
    def get_queryset(self):        		
        fclass_dict = dict(((fclass.classroom_id, fclass) for fclass in SpeculationClass.objects.filter(forum_id=self.kwargs['forum_id'])))		
        classroom_list = []
        classroom_ids = []
        classrooms = Classroom.objects.filter(teacher_id=self.request.user.id).order_by("-id")
        for classroom in classrooms:
            if classroom.id in fclass_dict:
                classroom_list.append([classroom, True, fclass_dict[classroom.id].deadline, fclass_dict[classroom.id].deadline_date])
            else :
                classroom_list.append([classroom, False, False, timezone.now()])
            classroom_ids.append(classroom.id)
        assistants = Assistant.objects.filter(user_id=self.request.user.id)
        for assistant in assistants:
            classroom = Classroom.objects.get(id=assistant.classroom_id)
            if not classroom.id in classroom_ids:
                if classroom.id in fclass_dict:
                    classroom_list.append([classroom, True, fclass_dict[classroom.id].deadline, fclass_dict[classroom.id].deadline_date])
                else :
                    classroom_list.append([classroom, False, False, timezone.now()])
        return classroom_list
			
    def get_context_data(self, **kwargs):
        context = super(SpeculationClassListView, self).get_context_data(**kwargs)				
        fwork = SpeculationWork.objects.get(id=self.kwargs['forum_id'])
        context['fwork'] = fwork
        context['forum_id'] = self.kwargs['forum_id']
        return context	
	
# Ajax 開放班取、關閉班級
def speculation_switch(request):
    forum_id = request.POST.get('forumid')
    classroom_id = request.POST.get('classroomid')		
    status = request.POST.get('status')
    try:
        fwork = SpeculationClass.objects.get(forum_id=forum_id, classroom_id=classroom_id)
        if status == 'false' :
            fwork.delete()
    except ObjectDoesNotExist:
        if status == 'true':
            fwork = SpeculationClass(forum_id=forum_id, classroom_id=classroom_id)
            fwork.save()
    return JsonResponse({'status':status}, safe=False)        
	
# 列出某作業所有同學名單
def speculation_class(request, classroom_id, work_id):
    enrolls = Enroll.objects.filter(classroom_id=classroom_id)
    classroom_name = Classroom.objects.get(id=classroom_id).name
    classmate_work = []
    scorer_name = ""
    for enroll in enrolls:
        try:    
            work = SWork.objects.get(student_id=enroll.student_id, index=work_id)
            if work.scorer > 0 :
                scorer = User.objects.get(id=work.scorer)
                scorer_name = scorer.first_name
            else :
                scorer_name = "1"
        except ObjectDoesNotExist:
            work = SWork(index=work_id, student_id=1)
        try:
            group_name = EnrollGroup.objects.get(id=enroll.group).name
        except ObjectDoesNotExist:
            group_name = "沒有組別"
        assistant = Assistant.objects.filter(classroom_id=classroom_id, student_id=enroll.student_id, lesson=work_id)
        if assistant.exists():
            classmate_work.append([enroll,work,1, scorer_name, group_name])
        else :
            classmate_work.append([enroll,work,0, scorer_name, group_name])   
    def getKey(custom):
        return custom[0].seat
	
    classmate_work = sorted(classmate_work, key=getKey)
   
    return render(request,'teacher/speculation_class.html',{'classmate_work': classmate_work, 'classroom_id':classroom_id, 'index': work_id})

# 列出所有思辨主題素材
class SpeculationContentListView(ListView):
    model = SpeculationContent
    context_object_name = 'contents'
    template_name = "teacher/speculation_content.html"		
    def get_queryset(self):
        queryset = SpeculationContent.objects.filter(forum_id=self.kwargs['forum_id'])
        return queryset
			
    def get_context_data(self, **kwargs):
        context = super(SpeculationContentListView, self).get_context_data(**kwargs)
        fwork = SpeculationWork.objects.get(id=self.kwargs['forum_id'])
        context['fwork']= fwork
        context['forum_id'] = self.kwargs['forum_id']
        return context	
			
#新增一個課程
class SpeculationContentCreateView(CreateView):
    model = SpeculationContent
    form_class = SpeculationContentForm
    template_name = "teacher/speculation_content_form.html"
    def form_valid(self, form):
        self.object = form.save(commit=False)
        work = SpeculationContent(forum_id=self.object.forum_id)
        if self.object.types == 1:
            work.types = 1
            work.link = self.object.link
            work.title = self.object.title            
        elif self.object.types  == 2:
            work.types = 2					
            work.youtube = self.object.youtube
        elif self.object.types  == 3:
            work.types = 3
            myfile = self.request.FILES['content_file']
            fs = FileSystemStorage()
            filename = uuid4().hex
            work.title = myfile.name
            work.filename = str(self.request.user.id)+"/"+filename
            fs.save("static/upload/"+str(self.request.user.id)+"/"+filename, myfile)	
        elif self.object.types  == 4:
            work.types = 4
            work.text = self.object.text
        work.memo = self.object.memo
        work.save()         

        #if self.object.types == 3:
        #  return JsonResponse({'files': [{'name': work.filename}]}, safe=False)
        return redirect("/teacher/speculation/content/"+str(self.kwargs['forum_id']))  

    def get_context_data(self, **kwargs):
        ctx = super(SpeculationContentCreateView, self).get_context_data(**kwargs)
        ctx['forum'] = SpeculationWork.objects.get(id=self.kwargs['forum_id'])
        return ctx

def speculation_delete(request, forum_id, content_id):
    instance = SpeculationContent.objects.get(id=content_id)
    instance.delete()

    return redirect("/teacher/speculation/content/"+str(forum_id))  
	
def speculation_edit(request, forum_id, content_id):
    try:
        instance = SpeculationContent.objects.get(id=content_id)
    except:
        pass
    if request.method == 'POST':
            content_id = request.POST.get("id", "")
            try:
                content = SpeculationContent.objects.get(id=content_id)
            except ObjectDoesNotExist:
	              content = SpeculattionContent(forum_id= request.POST.get("forum_id", ""), types=form.cleaned_data['types'])
            if content.types == 1:
                content.text = request.POST.get("text", "")
            elif content.types == 2:
                content.youtube = request.POST.get("youtube", "")
            elif content.types == 3:
                myfile =  request.FILES.get("content_file", "")
                fs = FileSystemStorage()
                filename = uuid4().hex
                content.title = myfile.name
                content.filename = str(request.user.id)+"/"+filename
                fs.save("static/upload/"+str(request.user.id)+"/"+filename, myfile)
            elif content.types == 4:
                content.title = request.POST.get("title", "")
                content.link = request.POST.get("link", "")								
            content.memo = request.POST.get("memo", "")
            content.save()
            return redirect('/teacher/speculation/content/'+forum_id)   
    return render(request,'teacher/speculation_edit.html',{'content': instance, 'forum_id':forum_id, 'content_id':content_id})		
	
def speculation_download(request, content_id):
    content = SpeculationContent.objects.get(id=content_id)
    filename = content.title
    BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    download =  BASE_DIR + "/static/upload/" + content.filename
    wrapper = FileWrapper(file( download, "r" ))
    response = HttpResponse(wrapper, content_type = 'application/force-download')
    #response = HttpResponse(content_type='application/force-download')
    response['Content-Disposition'] = 'attachment; filename={0}'.format(filename.encode('utf8'))
    # It's usually a good idea to set the 'Content-Length' header too.
    # You can also set any other required headers: Cache-Control, etc.
    return response
    #return render(request,'student/download.html', {'download':download})
		
class SpeculationEditUpdateView(UpdateView):
    model = SpeculationWork
    fields = ['title']
    template_name = 'form.html'
    #success_url = '/teacher/forum/domain/'
    def get_success_url(self):
        succ_url =  '/teacher/speculation/'+self.kwargs['classroom_id']
        return succ_url
	
def speculation_export(request, classroom_id, forum_id):
	if not is_teacher(request.user, classroom_id):
		return redirect("/")
	classroom = Classroom.objects.get(id=classroom_id)
	try:
		fwork = FWork.objects.get(id=forum_id)
		enrolls = Enroll.objects.filter(classroom_id=classroom_id)
		datas = []
		contents = FContent.objects.filter(forum_id=forum_id).order_by("-id")
		fwork = FWork.objects.get(id=forum_id)
		works_pool = SFWork.objects.filter(index=forum_id).order_by("-id")
		reply_pool = SFReply.objects.filter(index=forum_id).order_by("-id")	
		file_pool = SFContent.objects.filter(index=forum_id, visible=True).order_by("-id")	
		for enroll in enrolls:
			works = filter(lambda w: w.student_id==enroll.student_id, works_pool)
			if len(works)>0:
				replys = filter(lambda w: w.work_id==works[0].id, reply_pool)
			else:
				replys = []
			files = filter(lambda w: w.student_id==enroll.student_id, file_pool)
			if enroll.seat > 0:
				datas.append([enroll, works, replys, files])
		def getKey(custom):
			return -custom[0].seat
		datas = sorted(datas, key=getKey, reverse=True)	
		#word
		document = Document()
		docx_title=u"思辨區-" + classroom.name + "-"+ str(timezone.localtime(timezone.now()).date())+".docx"
		document.add_paragraph(request.user.first_name + u'的思辨區作業')
		document.add_paragraph(u'主題：'+fwork.title)		
		document.add_paragraph(u"班級：" + classroom.name)		
		
		for enroll, works, replys, files in datas:
			user = User.objects.get(id=enroll.student_id)
			run = document.add_paragraph().add_run(str(enroll.seat)+")"+user.first_name)
			font = run.font
			font.color.rgb = RGBColor(0xFA, 0x24, 0x00)
			if len(works)>0:
				#p = document.add_paragraph(str(works[0].publication_date)[:19]+'\n'+works[0].memo)
				p = document.add_paragraph(str(localtime(works[0].publication_date))[:19]+'\n')
				# 將 memo 以時間標記為切割點，切分為一堆 tokens
				tokens = re.split('(\[m_\d+#\d+:\d+:\d+\])', works[0].memo)
				# 依續比對 token 格式
				for token in tokens:
					m = re.match('\[m_(\d+)#(\d+):(\d+):(\d+)\]', token)
					if m: # 若為時間標記，則插入連結
						vid = filter(lambda material: material.id == int(m.group(1)), contents)[0]
						add_hyperlink(document, p, vid.youtube+"&t="+m.group(2)+"h"+m.group(3)+"m"+m.group(4)+"s", "["+m.group(2)+":"+m.group(3)+":"+m.group(4)+"]")
					else: # 以一般文字插入
						p.add_run(token)
			if len(replys)>0:
				for reply in replys:
					user = User.objects.get(id=reply.user_id)
					run = document.add_paragraph().add_run(user.first_name+u'>'+str(localtime(reply.publication_date))[:19]+u'>留言:\n'+reply.memo)
					font = run.font
					font.color.rgb = RGBColor(0x42, 0x24, 0xE9)		
			if len(files)>0:
				for file in files:
					if file.visible:
						if file.title[-3:].upper() == "PNG" or file.title[-3:].upper() == "JPG":
							filename = 'static/upload/'+file.filename
							if os.path.exists(filename):
								copyfile(filename, 'static/upload/file.png')					
								document.add_picture('static/upload/file.png',width=Inches(6.0))
						else:
							p = document.add_paragraph()
							full_url = request.build_absolute_uri()
							index = full_url.find("/",9)
							url = full_url[:index] + "/student/speculation/download/" + str(file.id) 
							add_hyperlink(document, p, url, file.title)
		# Prepare document for download        
		f = StringIO()
		document.save(f)
		length = f.tell()
		f.seek(0)
		response = HttpResponse(
			f.getvalue(),
			content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
		)
		response['Content-Disposition'] = 'attachment; filename={0}'.format(docx_title.encode('utf8')) 
		response['Content-Length'] = length
		return response

	except ObjectDoesNotExist:
		pass
	return True

def speculation_grade(request, classroom_id, action):
	classroom = Classroom.objects.get(id=classroom_id)
	forum_ids = []
	forums = []
	fclasses = SpeculationClass.objects.filter(classroom_id=classroom_id).order_by("publication_date", "forum_id")
	for fclass in fclasses:
		forum_ids.append(fclass.forum_id)
		forum = SpeculationWork.objects.get(id=fclass.forum_id)
		forums.append(forum.title)
	enrolls = Enroll.objects.filter(classroom_id=classroom_id).order_by("seat")
	datas = {}
	for enroll in enrolls:
			sfworks = SSpeculationWork.objects.filter(index__in=forum_ids, student_id=enroll.student_id).order_by("id")
			if len(sfworks) > 0:
				for fclass in fclasses:
						works = filter(lambda w: w.index==fclass.forum_id, sfworks)
						if enroll.student_id in datas:
							if len(works) > 0 :
								datas[enroll.student_id].append(works[0])
							else :
								datas[enroll.student_id].append(SSpeculationWork())
						else:
							if len(works) > 0:
								datas[enroll.student_id] = [works[0]]
							else :
								datas[enroll.student_id] = [SSpeculationWork()]
			else :
				datas[enroll.student_id] = [SSpeculationWork()]
	results = []
	for enroll in enrolls:
		student_name = User.objects.get(id=enroll.student_id).first_name
		results.append([enroll, student_name, datas[enroll.student_id]])
	
	#下載Excel
	if action == 1:
		classroom = Classroom.objects.get(id=classroom_id)       
		output = StringIO()
		workbook = xlsxwriter.Workbook(output)    
		worksheet = workbook.add_worksheet(classroom.name)
		date_format = workbook.add_format({'num_format': 'yy/mm/dd'})
		
		row = 1
		worksheet.write(row, 1, u'座號')
		worksheet.write(row, 2, u'姓名')
		index = 3
	
		for forum in forums:
			worksheet.write(row, index, forum)
			index += 1
		
		row += 1
		index = 3

		for fclass in fclasses:
			worksheet.write(row, index, datetime.strptime(str(fclass.publication_date)[:19],'%Y-%m-%d %H:%M:%S'), date_format)
			index += 1			

		for enroll, student_name, sworks in results:
			row += 1
			worksheet.write(row, 1, enroll.seat)
			worksheet.write(row, 2, student_name)
			index = 3
			for work in sworks:
				if work.id:
					worksheet.write(row, index, work.score)
				else:
					worksheet.write(row, index, '')
				index +=1 

		workbook.close()
		# xlsx_data contains the Excel file
		response = HttpResponse(content_type='application/vnd.ms-excel')
		filename = classroom.name + '-' + str(localtime(timezone.now()).date()) + '.xlsx'
		response['Content-Disposition'] = 'attachment; filename={0}'.format(filename.encode('utf8'))
		xlsx_data = output.getvalue()
		response.write(xlsx_data)
		return response
	else :
		return render(request,'teacher/speculation_grade.html',{'results':results, 'forums':forums, 'classroom_id':classroom_id, 'fclasses':fclasses})

def speculation_deadline(request, classroom_id, forum_id):
    forum = SpeculationWork.objects.get(id=forum_id)
    if request.method == 'POST':
        form = SpeculationCategroyForm(request.POST)
        if form.is_valid():
            forum.domains = request.POST.getlist('domains')
            forum.levels = request.POST.getlist('levels')	
            forum.save()
            return redirect('/teacher/speculation/'+classroom_id)
    else:
        fclass = SpeculationClass.objects.get(classroom_id=classroom_id, forum_id=forum_id)
        form = SpeculationDeadlineForm(instance=fclass)
    return render(request,'teacher/speculation_deadline_form.html',{'fclass':fclass})

	
# Ajax 設定期限、取消期限
def speculation_deadline_set(request):
    forum_id = request.POST.get('forumid')
    classroom_id = request.POST.get('classroomid')		
    status = request.POST.get('status')
    try:
        fclass = SpeculationClass.objects.get(forum_id=forum_id, classroom_id=classroom_id)
    except ObjectDoesNotExist:
        fclass = SpeculationClass(forum_id=forum_id, classroom_id=classroom_id)
    if status == 'True':
        fclass.deadline = True
    else :
        fclass.deadline = False
    fclass.save()
    return JsonResponse({'status':status}, safe=False)        

# Ajax 設定期限日期
def speculation_deadline_date(request):
    forum_id = request.POST.get('forumid')
    classroom_id = request.POST.get('classroomid')		
    deadline_date = request.POST.get('deadlinedate')
    try:
        fclass = SpeculationClass.objects.get(forum_id=forum_id, classroom_id=classroom_id)
    except ObjectDoesNotExist:
        fclass = SpeculationClass(forum_id=forum_id, classroom_id=classroom_id)
    #fclass.deadline_date = deadline_date.strftime('%d/%m/%Y')
    fclass.deadline_date = datetime.strptime(deadline_date, '%Y %B %d - %I:%M %p')
    fclass.save()
    return JsonResponse({'status':deadline_date}, safe=False)             
        
# 列出文字註記類別
class SpeculationAnnotationListView(ListView):
    model = SpeculationAnnotation
    context_object_name = 'contents'
    template_name = "teacher/speculation_annotation.html"		
    def get_queryset(self):
        queryset = SpeculationAnnotation.objects.filter(forum_id=self.kwargs['forum_id'])
        return queryset
			
    def get_context_data(self, **kwargs):
        context = super(SpeculationAnnotationListView, self).get_context_data(**kwargs)
        fwork = SpeculationWork.objects.get(id=self.kwargs['forum_id'])
        context['fwork']= fwork
        context['forum_id'] = self.kwargs['forum_id']
        return context	
			
#新增一個註記類別
class SpeculationAnnotationCreateView(CreateView):
    model = SpeculationAnnotation
    form_class = SpeculationAnnotationForm
    template_name = "teacher/speculation_annotation_form.html"
    def form_valid(self, form):
        self.object = form.save(commit=False)
        work = SpeculationAnnotation(forum_id=self.object.forum_id)
        work.kind = self.object.kind
        work.color = self.object.color
        work.save()         
  
        return redirect("/teacher/speculation/annotation/"+str(self.kwargs['forum_id']))  

    def get_context_data(self, **kwargs):
        ctx = super(SpeculationAnnotationCreateView, self).get_context_data(**kwargs)
        ctx['forum_id'] = self.kwargs['forum_id']
        return ctx

def speculation_annotation_delete(request, forum_id, content_id):
    instance = SpeculationAnnotation.objects.get(id=content_id)
    instance.delete()

    return redirect("/teacher/speculation/annotation/"+forum_id)  
	
def speculation_annotation_edit(request, forum_id, content_id):
    try:
        instance = SpeculationAnnotation.objects.get(id=content_id)
    except:
        pass
    if request.method == 'POST':
            content_id = request.POST.get("content_id", "")
            try:
                content = SpeculationAnnotation.objects.get(id=content_id)
            except ObjectDoesNotExist:
	              content = SpeculattionAnnotation(forum_id= request.POST.get("forum_id", ""))
            content.kind = request.POST.get("kind", "")
            content.color = request.POST.get("color", "")								
            content.save()
            return redirect('/teacher/speculation/annotation/'+forum_id)   
    return render(request,'teacher/speculation_annotation_form.html',{'content': instance, 'forum_id':forum_id, 'content_id':content_id})

def speculation_group(request, classroom_id, forum_id):
    title = SpeculationWork.objects.get(id=forum_id).title
    speculation = SpeculationClass.objects.get(forum_id=forum_id, classroom_id=classroom_id)
    groups = ClassroomGroup.objects.filter(classroom_id=speculation.classroom_id).order_by("-id")
    return render(request,'teacher/speculation_group.html',{'speculation': speculation, 'groups':groups, 'title':title})

def speculation_group_set(request):
    group_id = request.POST.get('groupid')
    forum_id = request.POST.get('forumid')
    classroom_id = request.POST.get('classroomid')
    if group_id and forum_id and classroom_id:      
        forum = SpeculationClass.objects.get(forum_id=forum_id, classroom_id=classroom_id)	
        if is_teacher(request.user, forum.classroom_id) or is_assistant(request.user, forum.classroom_id):
            forum.group = group_id
            forum.save()      
        return JsonResponse({'status':'ok'}, safe=False)
    else:
        return JsonResponse({'status':'fail'}, safe=False) 
# 列出所有教師
class TeacherListView(ListView):
    model = User
    context_object_name = 'teachers'
    template_name = 'teacher/member.html'
    paginate_by = 50
		
    def get_queryset(self):      
        teachers = Group.objects.get(name="teacher").user_set.all().order_by("-last_login")
        queryset = []
        classrooms = Classroom.objects.all()
        fworks = FWork.objects.all()
        sworks = SpeculationWork.objects.all()
        for teacher in teachers:
            rooms = filter(lambda w: w.teacher_id==teacher.id, classrooms)
            classroom_ids = []									
            for classroom in rooms:
                classroom_ids.append(classroom.id)
            enroll = Enroll.objects.filter(classroom_id__in=classroom_ids, seat__gt=0).count()
            fwork = filter(lambda w: w.teacher_id==teacher.id, fworks)
            swork = filter(lambda w: w.teacher_id==teacher.id, sworks)
            queryset.append([teacher, len(rooms), len(fwork), len(swork), enroll])
        return queryset
			

			
# 列出所有組別
class GroupListView(ListView):
    model = Group
    context_object_name = 'groups'
    template_name = 'teacher/group.html'
    paginate_by = 25
    def get_queryset(self):      
        queryset = ClassroomGroup.objects.filter(classroom_id=self.kwargs['classroom_id']).order_by("-id")
        return queryset
			
    def get_context_data(self, **kwargs):
        context = super(GroupListView, self).get_context_data(**kwargs)
        context['classroom_id'] = self.kwargs['classroom_id']
        return context				
			
#新增一個分組
class GroupCreateView(CreateView):
    model = ClassroomGroup
    form_class = GroupForm
    template_name = 'teacher/group_form.html'    
    def form_valid(self, form):
        self.object = form.save(commit=False)
        self.object.classroom_id = self.kwargs['classroom_id']
        if is_teacher(self.request.user, self.kwargs['classroom_id']) or is_assistant(self.request.user, self.kwargs['classroom_id']):
            self.object.save()
            # 隨機分組
            if form.cleaned_data['assign'] == 1:
                enrolls = Enroll.objects.filter(classroom_id=self.kwargs['classroom_id'], seat__gt=0).order_by('?')
                number = 0
                for enroll in enrolls:
                    group = StudentGroup(group_id=self.object.id, enroll_id=enroll.id, group=(number % self.object.numbers))
                    group.save()
                    number += 1
                self.object.opening=False
                self.object.save()
                
        return redirect("/student/group/list/"+ str(self.object.id))   
			
    def get_context_data(self, **kwargs):
        context = super(GroupCreateView, self).get_context_data(**kwargs)
        return context	
			
class GroupUpdateView(UpdateView):
    model = ClassroomGroup
    form_class = GroupForm2		
    template_name = 'form.html'
    def get_success_url(self):
        succ_url =  '/student/group/list/'+self.kwargs['pk']
        return succ_url
			
    def form_valid(self, form):
        if is_teacher(self.request.user, self.kwargs['classroom_id']) or is_assistant(self.request.user, self.kwargs['classroom_id']):
            group = ClassroomGroup.objects.get(id=self.kwargs['pk'])
            reduce = group.numbers - form.cleaned_data['numbers']
            if reduce > 0:
                for i in range(reduce):
                    StudentGroup.objects.filter(group_id=self.kwargs['pk'], group=group.numbers-i).delete()
            form.save()
        return HttpResponseRedirect(self.get_success_url())
			

# 分組
def make(request):
    group_id = request.POST.get('groupid')
    action = request.POST.get('action')
    if group_id and action :      
        group = ClassroomGroup.objects.get(id=group_id)	
        if is_teacher(request.user, group.classroom_id) or is_assistant(request.user, group.classroom_id):
            if action == 'open':            
                group.opening = True   
            else : 
                group.opening = False
            group.save()      
        return JsonResponse({'status':'ok'}, safe=False)
    else:
        return JsonResponse({'status':'fail'}, safe=False) 
		
			
# 列出所有測驗主題
class ExamListView(ListView):
    model = Exam
    context_object_name = 'exams'
    template_name = "teacher/exam_list.html"		
    paginate_by = 20
    def get_queryset(self):        
        exam_classes = ExamClass.objects.filter(classroom_id=self.kwargs['classroom_id']).order_by("-publication_date", "-exam_id")
        exams = []
        for exam_class in exam_classes:
            exam = Exam.objects.get(id=exam_class.exam_id)
            exams.append([exam, exam_class])
        return exams
			
    def get_context_data(self, **kwargs):
        context = super(ExamListView, self).get_context_data(**kwargs)
        classroom = Classroom.objects.get(id=self.kwargs['classroom_id'])
        context['classroom'] = classroom
        return context	
        
#新增一個測驗主題
class ExamCreateView(CreateView):
    model = Exam
    form_class = ExamForm
    template_name = "teacher/exam_form.html"
    def form_valid(self, form):
        self.object = form.save(commit=False)
        self.object.user_id = self.request.user.id
        self.object.classroom_id = self.kwargs['classroom_id']
        self.object.domains = self.request.POST.getlist('domains')
        self.object.levels = self.request.POST.getlist('levels')	        
        self.object.save()  
        classrooms = self.request.POST.getlist('classrooms')
        for classroom in classrooms:
          exam_class = ExamClass(exam_id=self.object.id, classroom_id=classroom)
          exam_class.save()
        
        return redirect("/teacher/exam/"+str(self.kwargs['classroom_id']))           
        
    def get_context_data(self, **kwargs):
        context = super(ExamCreateView, self).get_context_data(**kwargs)
        classroom_list = []
        classrooms = Classroom.objects.filter(teacher_id=self.request.user.id)
        for classroom in classrooms:
            classroom_list.append(classroom.id)
        assistants = Assistant.objects.filter(user_id=self.request.user.id)
        for assistant in assistants:
            if not assistant.classroom_id in classroom_list:
                classroom_list.append(assistant.classroom_id)
        classrooms = Classroom.objects.filter(id__in=classroom_list).order_by("-id")
        context['classrooms'] = classrooms
        context['classroom_id'] = self.kwargs['classroom_id']
        context['classroom'] = Classroom.objects.get(id=self.kwargs['classroom_id'])
        context['domains'] = Domain.objects.all()
        context['levels'] = Level.objects.all()
        return context	
  
        return redirect("/teacher/exam/"+str(self.kwargs['classroom_id'])) 			

def speculation_categroy(request, classroom_id, forum_id):
    forum = SpeculationWork.objects.get(id=forum_id)
    domains = Domain.objects.all()
    levels = Level.objects.all()		
    if request.method == 'POST':
        form = SpeculationCategroyForm(request.POST)
        if form.is_valid():
            forum.domains = request.POST.getlist('domains')
            forum.levels = request.POST.getlist('levels')	
            forum.save()
            return redirect('/teacher/speculation/'+classroom_id+'/#'+str(forum.id))
    else:
        form = SpeculationCategroyForm(instance=forum)
        
    return render(request,'teacher/speculation_categroy_form.html',{'domains': domains, 'levels':levels, 'classroom_id': classroom_id, 'speculation':forum})
	
def exam_categroy(request, classroom_id, exam_id):
    exam = Exam.objects.get(id=exam_id)
    domains = Domain.objects.all()
    levels = Level.objects.all()		
    if request.method == 'POST':
        form = ExamCategroyForm(request.POST)
        if form.is_valid():
            exam.domains = request.POST.getlist('domains')
            exam.levels = request.POST.getlist('levels')	
            exam.save()
            return redirect('/teacher/exam/'+classroom_id+'/#'+str(exam.id))
    else:
        form = ExamCategroyForm(instance=exam)
        
    return render(request,'teacher/exam_categroy_form.html',{'domains': domains, 'levels':levels, 'classroom_id': classroom_id, 'exam':exam})

	
# 列出所有討論主題
class ExamAllListView(ListView):
    model = Exam
    context_object_name = 'exams'
    template_name = "teacher/exam_all.html"		
    paginate_by = 20
		
    def get_queryset(self):
      # 年級
      if self.kwargs['categroy'] == 1:
        queryset = FWork.objects.filter(levels__contains=self.kwargs['categroy_id']).order_by("-id")
      # 學習領域
      elif self.kwargs['categroy'] == 2:
        queryset = FWork.objects.filter(domains__contains=self.kwargs['categroy_id']).order_by("-id")   
      else:
        queryset = FWork.objects.all().order_by("-id")
      if self.request.GET.get('account') != None:
        keyword = self.request.GET.get('account')
        users = User.objects.filter(Q(username__icontains=keyword) | Q(first_name__icontains=keyword)).order_by('-id')
        user_list = []
        for user in users:
            user_list.append(user.id)
        forums = queryset.filter(teacher_id__in=user_list)
        return forums
      else:				
        return queryset
			
    def get_context_data(self, **kwargs):
        context = super(ExamAllListView, self).get_context_data(**kwargs)
        context['categroy'] = self.kwargs['categroy']							
        context['categroy_id'] = self.kwargs['categroy_id']							
        context['levels'] = Level.objects.all()				
        context['domains'] = Domain.objects.all()
        return context	

# 列出某測驗主題的班級
class ExamClassListView(ListView):
    model = Exam
    context_object_name = 'classrooms'
    template_name = "teacher/exam_class.html"		
    paginate_by = 20
	
    def get_queryset(self):        		
        eclass_dict = dict(((eclass.classroom_id, eclass) for eclass in ExamClass.objects.filter(exam_id=self.kwargs['exam_id'])))		
        classroom_list = []
        classroom_ids = []
        classrooms = Classroom.objects.filter(teacher_id=self.request.user.id).order_by("-id")
        for classroom in classrooms:
            if classroom.id in eclass_dict:
                classroom_list.append([classroom, True, eclass_dict[classroom.id].deadline, eclass_dict[classroom.id].deadline_date])
            else :
                classroom_list.append([classroom, False, False, timezone.now()])
            classroom_ids.append(classroom.id)
        assistants = Assistant.objects.filter(user_id=self.request.user.id)
        for assistant in assistants:
            classroom = Classroom.objects.get(id=assistant.classroom_id)
            if not classroom.id in classroom_ids:
                if classroom.id in eclass_dict:
                    classroom_list.append([classroom, True, eclass_dict[classroom.id].deadline, eclass_dict[classroom.id].deadline_date])
                else :
                    classroom_list.append([classroom, False, False, timezone.now()])
        return classroom_list
			
    def get_context_data(self, **kwargs):
        context = super(ExamClassListView, self).get_context_data(**kwargs)				
        exam = Exam.objects.get(id=self.kwargs['exam_id'])
        context['exam'] = exam
        context['exam_id'] = self.kwargs['exam_id']
        return context	
	
# Ajax 開放班取、關閉班級
def exam_switch(request):
    exam_id = request.POST.get('examid')
    classroom_id = request.POST.get('classroomid')		
    status = request.POST.get('status')
    try:
        examclass = ExamClass.objects.get(exam_id=exam_id, classroom_id=classroom_id)
        if status == 'false' :
            examclass.delete()
    except ObjectDoesNotExist:
        if status == 'true':
            examclass = ExamClass(exam_id=exam_id, classroom_id=classroom_id)
            examclass.save()
    return JsonResponse({'status':status}, safe=False)        
	
class ExamEditUpdateView(UpdateView):
    model = Exam
    fields = ['title']
    template_name = 'form.html'
    #success_url = '/teacher/forum/domain/'
    def get_success_url(self):
        succ_url =  '/teacher/exam/'+str(self.kwargs['classroom_id'])
        return succ_url
		
def exam_deadline(request, classroom_id, exam_id):
    exam = Exam.objects.get(id=exam_id)
    if request.method == 'POST':
        form = ExamCategroyForm(request.POST)
        if form.is_valid():
            exam.domains = request.POST.getlist('domains')
            exam.levels = request.POST.getlist('levels')	
            forum.save()
            return redirect('/teacher/exam/'+classroom_id)
    else:
        examclass = ExamClass.objects.get(classroom_id=classroom_id, exam_id=exam_id)
        form = ExamDeadlineForm(instance=examclass)
    return render(request,'teacher/exam_deadline_form.html',{'examclass':examclass})

	
# Ajax 設定期限、取消期限
def exam_deadline_set(request):
    exam_id = request.POST.get('examid')
    classroom_id = request.POST.get('classroomid')		
    status = request.POST.get('status')
    try:
        examclass = ExamClass.objects.get(exam_id=exam_id, classroom_id=classroom_id)
    except ObjectDoesNotExist:
        examclass = Examclass(exam_id=exam_id, classroom_id=classroom_id)
    if status == 'True':
        examclass.deadline = True
    else :
        examclass.deadline = False
    examclass.save()
    return JsonResponse({'status':status}, safe=False)        

# Ajax 設定期限日期
def exam_deadline_date(request):
    exam_id = request.POST.get('examid')
    classroom_id = request.POST.get('classroomid')		
    deadline_date = request.POST.get('deadlinedate')
    try:
        examclass = ExamClass.objects.get(exam_id=exam_id, classroom_id=classroom_id)
    except ObjectDoesNotExist:
        examclass = ExamClass(exam_id=exam_id, classroom_id=classroom_id)
    #fclass.deadline_date = deadline_date.strftime('%d/%m/%Y')
    examclass.deadline_date = datetime.strptime(deadline_date, '%Y %B %d - %I:%M %p')
    examclass.save()
    return JsonResponse({'status':deadline_date}, safe=False)             
		
# 列出所有測驗題目
class ExamQuestionListView(ListView):
    model = ExamQuestion
    context_object_name = 'questions'
    template_name = "teacher/exam_question.html"		
    def get_queryset(self):
        queryset = ExamQuestion.objects.filter(exam_id=self.kwargs['exam_id']).order_by("-id")
        return queryset
			
    def get_context_data(self, **kwargs):
        context = super(ExamQuestionListView, self).get_context_data(**kwargs)
        exam = Exam.objects.get(id=self.kwargs['exam_id'])
        context['exam']= exam
        exam_classes = ExamClass.objects.filter(exam_id=self.kwargs['exam_id']).order_by("classroom_id")
        context['exam_classes'] = exam_classes
        questions = ExamQuestion.objects.filter(exam_id=self.kwargs['exam_id'])
        context['score_total'] = sum(question.score for question in questions)			
        return context	
			
#新增一個題目
class ExamQuestionCreateView(CreateView):
    model = ExamQuestion
    form_class = ExamQuestionForm
    template_name = "teacher/exam_question_form.html"
    def form_valid(self, form):
        self.object = form.save(commit=False)
        question = ExamQuestion(exam_id=self.object.exam_id, types=self.object.types)
        question.answer = self.object.answer
        if question.types == 2:
            question.option1 = self.object.option1	
            question.option2 = self.object.option2
            question.option3 = self.object.option3	
            question.option4 = self.object.option4        
           
        if 'title_pic' in self.request.FILES :
            myfile = self.request.FILES['title_pic']
            fs = FileSystemStorage()
            filename = uuid4().hex
            question.title_filename = str(self.request.user.id)+"/"+filename
            fs.save("static/exam/"+str(self.request.user.id)+"/"+filename, myfile)                
        question.title = self.object.title
        question.types = self.object.types
        question.score = self.object.score
        question.save()         
  
        return redirect("/teacher/exam/question/"+str(self.kwargs['exam_id']))  

    def get_context_data(self, **kwargs):
        ctx = super(ExamQuestionCreateView, self).get_context_data(**kwargs)
        ctx['exam'] = Exam.objects.get(id=self.kwargs['exam_id'])
        return ctx

def exam_check(request, exam_id, question_id):
    classroom_id = request.POST.get("class")
    q_answer = request.POST.get("answer")    
    if not is_teacher(request.user, classroom_id):
        return redirect("/")
    enroll_pool = Enroll.objects.filter(classroom_id=classroom_id).order_by("seat")		
    student_ids = map(lambda a: a.student_id, enroll_pool)
    answers =ExamAnswer.objects.filter(student_id__in=student_ids, question_id=question_id).order_by("student_id")	
    return render(request, 'teacher/exam_check.html', {'answers': answers, 'q_answer': q_answer, 'question_id':question_id})

# 分組
def exam_check_make(request):
    user_id = request.POST.get('userid')
    examwork_id = request.POST.get('examworkid')   
    question_id = request.POST.get('questionid')        
    action = request.POST.get('action')
    if user_id and examwork_id and question_id and action :
        try:  
            exam_answer = ExamAnswer.objects.get(student_id=user_id, examwork_id=examwork_id, question_id=question_id)            
            if action == 'set':
                exam_answer.answer_right = True 
            else :
                exam_answer.answer_right = False
            exam_answer.save()
        except ObjectDoesNotExist: 
            pass        
        try:
            examwork = ExamWork.objects.get(id=examwork_id)
            question_ids = [int(x) for x in examwork.questions.split(",")]
            questions = ExamQuestion.objects.filter(exam_id=examwork.exam_id)
            score = 0
            score_answer = dict((question.id, [question.score, question.answer]) for question in questions)			
            answer_dict = dict(((answer.question_id, [answer.answer, answer.answer_right]) for answer in ExamAnswer.objects.filter(examwork_id=examwork_id, question_id__in=question_ids, student_id=user_id)))		
            for question in questions:
                if question.id in answer_dict:
                    if score_answer[question.id][1] == answer_dict[question.id][0] or answer_dict[question.id][1]:
                       score += score_answer[question.id][0]
            examwork.score = score
            examwork.save()
        except ObjectDoesNotExist:
            pass
        return JsonResponse({'status':"ok"}, safe=False)
    else:
        return JsonResponse({'status':'fail'}, safe=False) 

def exam_publish(request, exam_id):
    exam = Exam.objects.get(id=exam_id)
    exam.opening = True
    exam.save()

    return redirect("/teacher/exam/question/"+str(exam_id))  

def exam_question_delete(request, exam_id, question_id):
    instance = ExamQuestion.objects.get(id=question_id)
    instance.delete()

    return redirect("/teacher/exam/question/"+str(exam_id))  
	
def exam_question_edit(request, exam_id, question_id):
    exam = Exam.objects.get(id=exam_id)
    try:
        instance = ExamQuestion.objects.get(id=question_id)
    except:
        pass
    if request.method == 'POST':
            question_id = request.POST.get("question_id")
            try:
                question = ExamQuestion.objects.get(id=question_id)
            except ObjectDoesNotExist:
	              question = ExamQuestion(exam_id= request.POST.get("exam_id"), types=form.cleaned_data['types'])
            if question.types == 2:
                question.option1 = request.POST.get("option1")	
                question.option2 = request.POST.get("option2")	
                question.option3 = request.POST.get("option3")	
                question.option4 = request.POST.get("option4")	
            question.score = request.POST.get("score")
            question.answer = request.POST.get("answer")	
            question.title = request.POST.get("title")
            if 'title_pic' in request.FILES :
                myfile = request.FILES['title_pic']
                fs = FileSystemStorage()
                filename = uuid4().hex
                question.title_filename = str(request.user.id)+"/"+filename
                fs.save("static/exam/"+str(request.user.id)+"/"+filename, myfile)  
            question.save()
            return redirect('/teacher/exam/question/'+exam_id+"#"+str(question.id))   
    return render(request,'teacher/exam_question_edit.html',{'question': instance, 'exam':exam, 'quesiton_id':question_id})	

# Create your views here.
def exam_import_sheet(request, types, exam_id):
    #if request.user.id != 1:
    #    return redirect("/")
    if request.method == "POST":
        form = UploadFileForm(request.POST,
                              request.FILES)
        if form.is_valid():
            ExamImportQuestion.objects.all().delete()            
            if types == "1":
                request.FILES['file'].save_to_database(
                    name_columns_by_row=0,
                    model=ExamImportQuestion,
                    mapdict=['title', 'answer', 'score'])
            elif types == "2":
                request.FILES['file'].save_to_database(
                    name_columns_by_row=0,
                    model=ExamImportQuestion,
                    mapdict=['title', 'option1', 'option2','option3','option4','answer', 'score'])                    
            elif types == "3":
                request.FILES['file'].save_to_database(
                    name_columns_by_row=0,
                    model=ExamImportQuestion,
                    mapdict=['title', 'answer', 'score'])
            questions = ExamImportQuestion.objects.all()
            return render(request, 'teacher/exam_import_question.html',{'types':types, 'questions':questions, 'exam_id': exam_id})
        else:
            return HttpResponseBadRequest()
    else:	
        form = UploadFileForm()
    return render(
        request,
        'teacher/exam_upload_form.html',
        {
					  'exam_id': exam_id, 
            'form': form,
            'title': 'Excel file upload and download example',
            'header': ('Please choose any excel file ' +
                       'from your cloned repository:')
        })
	
# Create your views here.
def exam_import_question(request, types, exam_id):
    #if request.user.id != 1:
    #    return redirect("/")
    questions = ExamImportQuestion.objects.all()           
    if types == "1":
        for question in questions:            
            new_question = ExamQuestion(exam_id=exam_id, types=1, title=question.title, answer=question.answer, score=question.score)        
            new_question.save()
    elif types == "2":
        for question in questions:         
            new_question = ExamQuestion(exam_id=exam_id, types=2, title=question.title, option1=question.option1, option2=question.option2, option3=question.option3, option4=question.option4, answer=question.answer, score=question.score)
            new_question.save()
    elif types == "3":
        for question in questions:     
            new_question = ExamQuestion(exam_id=exam_id, types=3, title=question.title, answer=question.answer, score=question.score)
            new_question.save()
            
    return redirect('/teacher/exam/question/'+exam_id)			
	
def exam_round(request, classroom_id, exam_id):
    if not is_teacher(request.user, classroom_id):
        return redirect("/")
    examclass = ExamClass.objects.get(classroom_id=classroom_id, exam_id=exam_id)
    return render(request,'teacher/exam_round.html',{'examclass':examclass})		
	
def exam_round_set(request):
    exam_id = request.POST.get('examid')
    classroom_id = request.POST.get('classroomid')		
    round_limit = request.POST.get('round_limit')
    try:
        examclass = ExamClass.objects.get(exam_id=exam_id, classroom_id=classroom_id)
    except ObjectDoesNotExist:
        examclass = Examclass(exam_id=exam_id, classroom_id=classroom_id)
    examclass.round_limit = int(round_limit)
    examclass.save()
    return JsonResponse({'status':'ok'}, safe=False)  	
	
def exam_score(request, classroom_id, exam_id):
    if not is_teacher(request.user, classroom_id):
        return redirect("/")
    exam = Exam.objects.get(id=exam_id)
    classroom = Classroom.objects.get(id=classroom_id)
    examclass = ExamClass.objects.get(classroom_id=classroom_id, exam_id=exam_id)
    enrolls = Enroll.objects.filter(classroom_id=classroom_id).order_by("seat")
    enroll_ids = []
    for enroll in enrolls:
        enroll_ids.append(enroll.student_id)
    examworks = ExamWork.objects.filter(exam_id=exam_id, student_id__in=enroll_ids, publish=True).order_by("-id")
    scores = []
    for enroll in enrolls:
        works = list(filter(lambda w: w.student_id == enroll.student_id, examworks))
        if len(works) > 0 :
            score_max = max(work.score for work in works)
            score_avg = sum(work.score for work in works) / len(works)	
        else :
            score_max = 0
            score_avg = 0
        scores.append([enroll, works, score_avg, score_max])
    return render(request,'teacher/exam_score.html',{'classroom': classroom, 'exam':exam, 'scores':scores})		
	
def exam_excel(request, classroom_id, exam_id):
    if not is_teacher(request.user, classroom_id):
        return redirect("/")

    exam = Exam.objects.get(id=exam_id)
    classroom = Classroom.objects.get(id=classroom_id)
    examclass = ExamClass.objects.get(classroom_id=classroom_id, exam_id=exam_id)
    enrolls = Enroll.objects.filter(classroom_id=classroom_id).order_by("seat")
    enroll_ids = []
    for enroll in enrolls:
        enroll_ids.append(enroll.student_id)
    examworks = ExamWork.objects.filter(exam_id=exam_id, student_id__in=enroll_ids, publish=True).order_by("-id")
    scores = []
    for enroll in enrolls:
        works = list(filter(lambda w: w.student_id == enroll.student_id, examworks))
        if len(works) > 0 :
            score_max = max(work.score for work in works)
            score_avg = sum(work.score for work in works) / len(works)	
        else :
            score_max = 0
            score_avg = 0
        scores.append([enroll, works, score_avg, score_max])

    claassroom = Classroom.objects.get(id=classroom_id)       
    output = io.BytesIO()
    workbook = xlsxwriter.Workbook(output)    
    worksheet = workbook.add_worksheet(classroom.name)
    date_format = workbook.add_format({'num_format': 'yy/mm/dd'})
		
    row = 1
    worksheet.write(row, 1, u'座號')
    worksheet.write(row, 2, u'姓名')
    worksheet.write(row, 3, u'次數')
    worksheet.write(row, 4, u'平均')
    worksheet.write(row, 5, u'最高分')
    worksheet.write(row, 6, u'分數')             

    for enroll, works, score_avg, acore_max in scores:
        row += 1
        worksheet.write(row, 1, enroll.seat)
        worksheet.write(row, 2, enroll.student.first_name)
        worksheet.write(row, 3, len(works))
        worksheet.write(row, 4, score_avg)
        worksheet.write(row, 5, score_max)
        index = 5
        for work in works:
            index += 1
            worksheet.write(row, index, work.score)              

    workbook.close()
	# xlsx_data contains the Excel file
    response = HttpResponse(content_type='application/vnd.ms-excel')
    filename = classroom.name + '-' + exam.title + '-' + str(localtime(timezone.now()).date()) + '.xlsx'
    filename_header = filename_browser(request, filename)
    response['Content-Disposition'] = 'attachment; ' + filename_header	
    xlsx_data = output.getvalue()
    response.write(xlsx_data)
    return response


# 列出所有討論測驗
class ExamAllListView(ListView):
    model = Exam
    context_object_name = 'exams'
    template_name = "teacher/exam_all.html"		
    paginate_by = 20
		
    def get_queryset(self):
      # 年級
      if self.kwargs['categroy'] == 1:
        queryset = Exam.objects.filter(levels__contains=self.kwargs['categroy_id']).order_by("-id")
      # 學習領域
      elif self.kwargs['categroy'] == 2:
        queryset = Exam.objects.filter(domains__contains=self.kwargs['categroy_id']).order_by("-id")   
      else:
        queryset = Exam .objects.all().order_by("-id")
      if self.request.GET.get('account') != None:
        keyword = self.request.GET.get('account')
        users = User.objects.filter(Q(username__icontains=keyword) | Q(first_name__icontains=keyword)).order_by('-id')
        user_list = []
        for user in users:
            user_list.append(user.id)
        exams = queryset.filter(teacher_id__in=user_list)
        return exams
      else:				
        return queryset
			
    def get_context_data(self, **kwargs):
        context = super(ExamAllListView, self).get_context_data(**kwargs)
        context['categroy'] = self.kwargs['categroy']							
        context['categroy_id'] = self.kwargs['categroy_id']							
        context['levels'] = Level.objects.all()				
        context['domains'] = Domain.objects.all()
        return context	

       

# 列出所有討論主題
class TeamListView(ListView):
    model = TeamWork
    context_object_name = 'teams'
    template_name = "teacher/team_list.html"		
    paginate_by = 20
    def get_queryset(self):        
        teamclasses = TeamClass.objects.filter(classroom_id=self.kwargs['classroom_id']).order_by("-publication_date", "-team_id")
        teams = []
        for teamclass in teamclasses:
            team = TeamWork.objects.get(id=teamclass.team_id)
            teams.append([team, teamclass])
        return teams
			
    def get_context_data(self, **kwargs):
        context = super(TeamListView, self).get_context_data(**kwargs)
        classroom = Classroom.objects.get(id=self.kwargs['classroom_id'])
        context['classroom'] = classroom
        return context	
        
#新增一個討論主題
class TeamCreateView(CreateView):
    model = TeamWork
    form_class = TeamForm
    template_name = "teacher/team_form.html"
    def form_valid(self, form):
        self.object = form.save(commit=False)
        self.object.teacher_id = self.request.user.id
        self.object.classroom_id = self.kwargs['classroom_id']
        self.object.domains = self.request.POST.getlist('domains')
        self.object.levels = self.request.POST.getlist('levels')	        
        self.object.save()  
        classrooms = self.request.POST.getlist('classrooms')
        for classroom in classrooms:
          teamclass = TeamClass(team_id=self.object.id, classroom_id=classroom)
          teamclass.save()
        
        return redirect("/teacher/team/"+str(self.kwargs['classroom_id']))           
        
    def get_context_data(self, **kwargs):
        context = super(TeamCreateView, self).get_context_data(**kwargs)
        classroom_list = []
        classrooms = Classroom.objects.filter(teacher_id=self.request.user.id)
        for classroom in classrooms:
            classroom_list.append(classroom.id)
        assistants = Assistant.objects.filter(user_id=self.request.user.id)
        for assistant in assistants:
            if not assistant.classroom_id in classroom_list:
                classroom_list.append(assistant.classroom_id)
        classrooms = Classroom.objects.filter(id__in=classroom_list).order_by("-id")
        context['classrooms'] = classrooms
        context['classroom_id'] = int(self.kwargs['classroom_id'])
        context['classroom'] = Classroom.objects.get(id=self.kwargs['classroom_id'])
        context['domains'] = Domain.objects.all()
        context['levels'] = Level.objects.all()
        return context	
	
			
def team_categroy(request, classroom_id, team_id):
    team = TeamWork.objects.get(id=team_id)
    domains = Domain.objects.all()
    levels = Level.objects.all()		
    if request.method == 'POST':
        form = TeamCategroyForm(request.POST)
        if form.is_valid():
            team.domains = request.POST.getlist('domains')
            team.levels = request.POST.getlist('levels')	
            team.save()
            return redirect('/teacher/team/'+classroom_id+'/#'+str(team.id))
    else:
        form = TeamCategroyForm(instance=team)
        
    return render(request,'teacher/team_categroy_form.html',{'domains': domains, 'levels':levels, 'classroom_id': classroom_id, 'team':team})


# 列出某任務主題的班級
class TeamClassListView(ListView):
    model = TeamWork
    context_object_name = 'classrooms'
    template_name = "teacher/team_class.html"		
    paginate_by = 20
	
    def get_queryset(self):        		
        teamclass_dict = dict(((teamclass.classroom_id, teamclass) for teamclass in TeamClass.objects.filter(team_id=self.kwargs['team_id'])))		
        classroom_list = []
        classroom_ids = []
        classrooms = Classroom.objects.filter(teacher_id=self.request.user.id).order_by("-id")
        for classroom in classrooms:
            if classroom.id in teamclass_dict:
                classroom_list.append([classroom, True, teamclass_dict[classroom.id].deadline, teamclass_dict[classroom.id].deadline_date])
            else :
                classroom_list.append([classroom, False, False, timezone.now()])
            classroom_ids.append(classroom.id)
        assistants = Assistant.objects.filter(user_id=self.request.user.id)
        for assistant in assistants:
            classroom = Classroom.objects.get(id=assistant.classroom_id)
            if not classroom.id in classroom_ids:
                if classroom.id in teamclass_dict:
                    classroom_list.append([classroom, True, teamclass_dict[classroom.id].deadline, teamclass_dict[classroom.id].deadline_date])
                else :
                    classroom_list.append([classroom, False, False, timezone.now()])
        return classroom_list
			
    def get_context_data(self, **kwargs):
        context = super(TeamClassListView, self).get_context_data(**kwargs)				
        teamwork = TeamWork.objects.get(id=self.kwargs['team_id'])
        context['teamwork'] = teamwork
        context['team_id'] = self.kwargs['team_id']
        return context	
	
# Ajax 開放班取、關閉班級
def team_switch(request):
    team_id = request.POST.get('teamid')
    classroom_id = request.POST.get('classroomid')		
    status = request.POST.get('status')
    try:
        teamwork = TeamClass.objects.get(team_id=team_id, classroom_id=classroom_id)
        if status == 'false' :
            teamwork.delete()
    except ObjectDoesNotExist:
        if status == 'true':
            teamwork = TeamClass(team_id=team_id, classroom_id=classroom_id)
            teamwork.save()
    return JsonResponse({'status':status}, safe=False)        
	
# 列出某任務所有同學名單
def team_class(request, classroom_id, work_id):
    enrolls = Enroll.objects.filter(classroom_id=classroom_id)
    classroom_name = Classroom.objects.get(id=classroom_id).name
    classmate_work = []
    scorer_name = ""
    for enroll in enrolls:
        try:    
            work = SWork.objects.get(student_id=enroll.student_id, index=work_id)
            if work.scorer > 0 :
                scorer = User.objects.get(id=work.scorer)
                scorer_name = scorer.first_name
            else :
                scorer_name = "1"
        except ObjectDoesNotExist:
            work = SWork(index=work_id, student_id=1)
        try:
            group_name = EnrollGroup.objects.get(id=enroll.group).name
        except ObjectDoesNotExist:
            group_name = "沒有組別"
        assistant = Assistant.objects.filter(classroom_id=classroom_id, student_id=enroll.student_id, lesson=work_id)
        if assistant.exists():
            classmate_work.append([enroll,work,1, scorer_name, group_name])
        else :
            classmate_work.append([enroll,work,0, scorer_name, group_name])   
    def getKey(custom):
        return custom[0].seat
	
    classmate_work = sorted(classmate_work, key=getKey)
   
    return render(request,'teacher/twork_class.html',{'classmate_work': classmate_work, 'classroom_id':classroom_id, 'index': work_id})
			
def team_deadline(request, classroom_id, team_id):
    team = TeamWork.objects.get(id=team_id)
    classroom = Classroom.objects.get(id=classroom_id)
    if request.method == 'POST':
        form = CategroyForm(request.POST)
        if form.is_valid():
            team.domains = request.POST.getlist('domains')
            team.levels = request.POST.getlist('levels')	
            team.save()
            return redirect('/teacher/team/'+classroom_id)
    else:
        teamclass = TeamClass.objects.get(classroom_id=classroom_id, team_id=team_id)
        form = TeamDeadlineForm(instance=teamclass)
        teamclasses = TeamClass.objects.filter(team_id=team_id).order_by("-id")
    return render(request,'teacher/team_deadline_form.html',{'teamclasses':teamclasses, 'teamclass':teamclass, 'team':team, 'classroom':classroom})

# Ajax 設定期限、取消期限
def team_deadline_set(request):
    team_id = request.POST.get('teamid')
    classroom_id = request.POST.get('classroomid')		
    status = request.POST.get('status')
    try:
        teamclass = TeamClass.objects.get(team_id=team_id, classroom_id=classroom_id)
    except ObjectDoesNotExist:
        teamclass = Teamclass(team_id=team_id, classroom_id=classroom_id)
    if status == 'True':
        teamclass.deadline = True
    else :
        teamclass.deadline = False
    teamclass.save()
    return JsonResponse({'status':status}, safe=False)        

# Ajax 設定期限日期
def team_deadline_date(request):
    team_id = request.POST.get('teamid')
    classroom_id = request.POST.get('classroomid')		
    deadline_date = request.POST.get('deadlinedate')
    try:
        teamclass = TeamClass.objects.get(team_id=team_id, classroom_id=classroom_id)
    except ObjectDoesNotExist:
        teamclass = TeamClass(team_id=team_id, classroom_id=classroom_id)
    #fclass.deadline_date = deadline_date.strftime('%d/%m/%Y')
    teamclass.deadline_date = datetime.strptime(deadline_date, '%Y %B %d - %H:%M')
    teamclass.save()
    return JsonResponse({'status':deadline_date}, safe=False)            		
	
			
class TeamEditUpdateView(UpdateView):
    model = TeamWork
    fields = ['title']
    template_name = 'form.html'
    #success_url = '/teacher/forum/domain/'
    def get_success_url(self):
        succ_url =  '/teacher/team/'+str(self.kwargs['classroom_id'])
        return succ_url
			
def team_group(request, classroom_id, team_id):
    enrolls = Enroll.objects.filter(classroom_id=classroom_id)
    enroll_dict = {}
    for enroll in enrolls:
        enroll_dict[enroll.id] = enroll
    groups = ClassroomGroup.objects.filter(classroom_id=classroom_id)
    group_ids = []
    for group in groups:
        group_ids.append(group.id)
    classroom = Classroom.objects.get(id=classroom_id)
    studentgroups = StudentGroup.objects.filter(group_id__in=group_ids)
    group_list = []
    for group in groups:        
        groupclass_list = []  
        groupclass_dict = {}
        students = filter(lambda student: student.group_id == group.id , studentgroups)
        for student in students:
            if student.enroll_id in enroll_dict:
                if student.group in groupclass_dict:
                    groupclass_dict[student.group].append(enroll_dict[student.enroll_id])
                else :
                    groupclass_dict[student.group] = [enroll_dict[student.enroll_id]]
        for key in groupclass_dict:
            groupclass_list.append([key, groupclass_dict[key]])
        group_list.append([group.id, groupclass_list])
    teamclass = TeamClass.objects.get(team_id=team_id, classroom_id=classroom_id)
    try:
        group = ClassroomGroup.objects.get(id=teamclass.group)
    except ObjectDoesNotExist:
        group = ClassroomGroup(title="不分組", id=0)
    return render(request,'teacher/team_group.html',{'team_id': team_id, 'teamgroup': group, 'groups':groups, 'classroom':classroom, 'group_list':group_list})

# 影片觀看時間統計
class EventVideoView(ListView):
    context_object_name = 'events'
    #paginate_by = 50
    template_name = 'teacher/event_video.html'

    def get_queryset(self):    
                enrolls = Enroll.objects.filter(classroom_id=self.kwargs['classroom_id'], seat__gt=0).order_by("seat")
                events = []
                for enroll in enrolls: 
                        videos = VideoLogHelper().getLogByUserid(enroll.student_id,self.kwargs['work_id'])
                        length = 0
                        for video in videos: 
                            length += video['length']
                        events.append([enroll, length/60.0])
                return events
			
    def get_context_data(self, **kwargs):
        context = super(EventVideoView, self).get_context_data(**kwargs)
        classroom = Classroom.objects.get(id=self.kwargs['classroom_id'])
        context['length'] = FContent.objects.get(id=self.kwargs['work_id']).youtube_length / 60.0
        context['content_id'] = self.kwargs['work_id']
        context['classroom'] = classroom
        enrolls = Enroll.objects.filter(classroom_id=classroom.id)
        context['height'] = 100 + enrolls.count() * 40
        return context
			
# 記錄影片長度
def video_length(request):
    content_id = request.POST.get('content_id')
    length = request.POST.get('length')
    page = request.POST.get('page')
    if page == "team":
        teamcontent = TeamContent.objects.get(id=content_id)
        teamcontent.youtube_length = length
        teamcontent.save()
    elif page == "forum":
        fcontent = FContent.objects.get(id=content_id)
        fcontent.youtube_length = length
        fcontent.save()
    elif page == "couurse":
        coursecontent = CourseContent.objects.get(id=content_id)
        coursecontent.youtube_length = length
        coursecontent.save()        
    return JsonResponse({'status':'ok'}, safe=False)	
	
# 影片記錄條
class VideoListView(ListView):
    context_object_name = 'videos'
    template_name = 'teacher/event_video_user.html'
    
    def get_queryset(self):
        videos = VideoLogHelper().getLogByUserid(self.kwargs['user_id'],self.kwargs['content_id'])        
        return videos
        
    def get_context_data(self, **kwargs):
        context = super(VideoListView, self).get_context_data(**kwargs)
        content = FContent.objects.get(id=self.kwargs['content_id'])
        context['user_id'] = self.kwargs['user_id']
        context['content'] = content
        context['length'] = content.youtube_length
        return context  

    # 限本班任課教師或助教     
    def render_to_response(request,self, context):
        if not is_teacher(self.request.user ,self.kwargs['classroom_id']):
            if not is_assistant(self.request.user, self.kwargs['classroom_id'] ):
                  return redirect('/')
        return super(VideoListView, self).render(request,context)   
			
# Ajax 設定合作區組別
def team_group_set(request):
    team_id = request.POST.get('teamid')
    classroom_id = request.POST.get('classroomid')		
    group = request.POST.get('groupid')
    try:
        teamclass = TeamClass.objects.get(team_id=int(team_id), classroom_id=int(classroom_id))
    except ObjectDoesNotExist:
        pass
    teamclass.group = int(group)
    teamclass.save()    
    return JsonResponse({'status':team_id}, safe=False)  
